"""Script para actualizar Documentacion_ERP_Ciete.docx con los nuevos módulos."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

PATH = 'C:/Users/Bernardo/Documents/erp-ciete/erp-ciete/erp-ciete/Documentacion_ERP_Ciete.docx'
doc = Document(PATH)

# ─────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────

def replace_para_text(para, new_text):
    """Sustituye todo el texto de un párrafo conservando el estilo."""
    # Guardar propiedades del primer run (fuente, etc.)
    first_runs = para._element.findall(qn('w:r'))
    rPr = None
    if first_runs:
        rPr_elem = first_runs[0].find(qn('w:rPr'))
        if rPr_elem is not None:
            rPr = deepcopy(rPr_elem)
    # Eliminar todos los runs existentes
    for r in para._element.findall(qn('w:r')):
        para._element.remove(r)
    # Añadir un único run con el nuevo texto
    new_r = OxmlElement('w:r')
    if rPr is not None:
        new_r.append(rPr)
    new_t = OxmlElement('w:t')
    new_t.text = new_text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    para._element.append(new_r)


def clone_para_with_text(source_para, new_text):
    """Clona un párrafo (conservando su estilo/sangría) con nuevo texto."""
    new_p = deepcopy(source_para._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    new_r = OxmlElement('w:r')
    # Copiar propiedades del primer run original si existen
    for orig_r in source_para._element.findall(qn('w:r')):
        rPr = orig_r.find(qn('w:rPr'))
        if rPr is not None:
            new_r.append(deepcopy(rPr))
        break
    new_t = OxmlElement('w:t')
    new_t.text = new_text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p


def detach_and_return(elem):
    """Desconecta un elemento XML de su padre y lo devuelve."""
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)
    return elem


def make_db_section(table_name, fields):
    """Crea heading H2 + tabla de BD; devuelve (h_elem, tbl_elem) desconectados."""
    h = doc.add_paragraph(style='Heading 2')
    h.clear()
    h.add_run(f'Tabla: {table_name}')
    h_elem = detach_and_return(h._element)

    t = doc.add_table(rows=1, cols=4, style='LightGrid-Accent1')
    for i, txt in enumerate(['Campo', 'Tipo', 'Nulo', 'Descripción']):
        t.rows[0].cells[i].text = txt
    for field in fields:
        row = t.add_row()
        for i, val in enumerate(field):
            row.cells[i].text = val
    tbl_elem = detach_and_return(t._tbl)

    return h_elem, tbl_elem


# ─────────────────────────────────────────────────────────
# 1. VERSIÓN Y FECHA
# ─────────────────────────────────────────────────────────
for p in doc.paragraphs:
    if 'Versión 2.0' in p.text and '16/04/2026' in p.text:
        for run in p.runs:
            run.text = (run.text
                        .replace('Versión 2.0', 'Versión 3.0')
                        .replace('16/04/2026', '25/04/2026'))
        print('✓ Versión actualizada a 3.0 / 25/04/2026')

# ─────────────────────────────────────────────────────────
# 2. ÍNDICE — añadir entradas 9-12 después del punto 8
# ─────────────────────────────────────────────────────────
toc_ref = None
for p in doc.paragraphs:
    if 'Módulo de Estaciones de Servicio' in p.text and p.style.name == 'Normal':
        toc_ref = p
        break

if toc_ref:
    new_toc_items = [
        '  9.   Módulo de Usuarios',
        '  10.  Módulo de Contratos',
        '  11.  Módulo de Tarifario',
        '  12.  Módulo de Facturas',
    ]
    last = toc_ref._element
    for item in new_toc_items:
        new_p = clone_para_with_text(toc_ref, item)
        last.addnext(new_p)
        last = new_p
    print('✓ Índice actualizado con 4 nuevas entradas')

# ─────────────────────────────────────────────────────────
# 3. ESTRUCTURA DE CARPETAS — Backend (párrafo 21)
# ─────────────────────────────────────────────────────────
NEW_BACKEND = (
    'backend/\n'
    '├── main.py                  ← Punto de entrada FastAPI. Configura CORS,\n'
    '│                              registra routers, ejecuta migración de BD\n'
    '│                              y crea el usuario admin al arrancar.\n'
    '├── database.py              ← Conexión SQLite con SQLAlchemy.\n'
    '│                              Expone get_db() como dependencia.\n'
    '├── auth.py                  ← JWT (python-jose) + hashing bcrypt.\n'
    '│                              Funciones: create_access_token,\n'
    '│                              get_current_user, verify_password,\n'
    '│                              get_password_hash.\n'
    '├── requirements.txt         ← Dependencias Python del proyecto.\n'
    '├── database.sqlite          ← Archivo de base de datos (auto-generado).\n'
    '│\n'
    '├── models/                  ← Modelos SQLAlchemy (tablas de la BD)\n'
    '│   ├── user.py              ← Tabla users (+ campo permisos)\n'
    '│   ├── presupuesto.py       ← Tablas presupuestos + presupuestos_lineas\n'
    '│   ├── pedido.py            ← Tablas pedidos + lineas_pedido\n'
    '│   ├── empresa.py           ← Tablas empresas + contactos_empresas\n'
    '│   ├── contacto.py          ← Tabla contactos\n'
    '│   ├── estacion.py          ← Tabla es (estaciones de servicio)\n'
    '│   ├── contrato.py          ← Tabla contratos                    ← NUEVO\n'
    '│   ├── tarifario.py         ← Tablas tarifario + tarifario_servicios  ← NUEVO\n'
    '│   └── factura.py           ← Tablas facturas + factura_pedidos\n'
    '│                              + lineas_factura + cobros           ← NUEVO\n'
    '│\n'
    '├── schemas/                 ← Schemas Pydantic (validación + serialización)\n'
    '│   ├── user.py              ← UserOut, UserCreate, Token, LoginRequest\n'
    '│   ├── presupuesto.py       ← PresupuestoOut, PresupuestoCreate, etc.\n'
    '│   ├── pedido.py            ← PedidoOut, LineaPedidoOut\n'
    '│   ├── empresa.py           ← EmpresaOut, EmpresaCreate, etc.\n'
    '│   ├── contacto.py          ← ContactoOut, ContactoCreate\n'
    '│   ├── estacion.py          ← EstacionOut, EstacionCreate\n'
    '│   ├── usuario.py           ← UsuarioOut, UsuarioCreate, UsuarioUpdate  ← NUEVO\n'
    '│   ├── contrato.py          ← ContratoOut, ContratoCreate               ← NUEVO\n'
    '│   ├── tarifario.py         ← TarifarioOut, TarifarioServicioOut, etc.  ← NUEVO\n'
    '│   └── factura.py           ← FacturaOut, CobroOut, LineaFacturaOut      ← NUEVO\n'
    '│\n'
    '└── routers/                 ← Endpoints REST organizados por área\n'
    '    ├── auth.py              ← POST /auth/login, /auth/registro, GET /auth/me\n'
    '    ├── presupuestos.py      ← CRUD presupuestos + cambio estado + convertir\n'
    '    ├── pedidos.py           ← GET /pedidos, GET /pedidos/{id}\n'
    '    ├── empresas.py          ← CRUD empresas + gestión de contactos\n'
    '    ├── contactos.py         ← CRUD contactos\n'
    '    ├── estaciones.py        ← CRUD estaciones de servicio\n'
    '    ├── contratos.py         ← CRUD contratos                     ← NUEVO\n'
    '    ├── tarifario.py         ← CRUD tarifarios + gestión servicios ← NUEVO\n'
    '    ├── facturas.py          ← CRUD facturas + cobros              ← NUEVO\n'
    '    └── usuarios.py          ← CRUD usuarios (solo admin)          ← NUEVO'
)

para21 = doc.paragraphs[21]
replace_para_text(para21, NEW_BACKEND)
print('✓ Estructura de carpetas backend actualizada')

# ─────────────────────────────────────────────────────────
# 4. ESTRUCTURA DE CARPETAS — Frontend (párrafo 26)
# ─────────────────────────────────────────────────────────
NEW_FRONTEND = (
    'frontend/src/app/\n'
    '├── app.component.ts         ← Componente raíz (solo renderiza <router-outlet>)\n'
    '├── app.config.ts            ← Configura provideHttpClient + provideRouter\n'
    '├── app.routes.ts            ← Definición de rutas con lazy loading\n'
    '│\n'
    '├── core/\n'
    '│   ├── services/\n'
    '│   │   ├── auth.service.ts          ← login, logout, getToken, getMe\n'
    '│   │   ├── presupuestos.service.ts  ← getAll, getOne, create,\n'
    '│   │   │                               cambiarEstado, convertir\n'
    '│   │   ├── pedidos.service.ts       ← getAll, getOne\n'
    '│   │   ├── empresas.service.ts      ← getAll, getOne, create, update, delete\n'
    '│   │   ├── contactos.service.ts     ← getAll, getOne, create, update, delete\n'
    '│   │   ├── estaciones.service.ts    ← getAll, getOne, create, update, delete\n'
    '│   │   ├── usuarios.service.ts      ← CRUD usuarios                  ← NUEVO\n'
    '│   │   ├── contratos.service.ts     ← CRUD contratos                 ← NUEVO\n'
    '│   │   ├── tarifario.service.ts     ← CRUD tarifarios + servicios    ← NUEVO\n'
    '│   │   └── facturas.service.ts      ← CRUD facturas + cobros         ← NUEVO\n'
    '│   ├── interceptors/\n'
    '│   │   └── auth.interceptor.ts      ← Añade Bearer token a cada petición\n'
    '│   └── guards/\n'
    '│       ├── auth.guard.ts            ← Redirige a /login si no hay token\n'
    '│       └── permisos.guard.ts        ← Comprueba permisos por módulo  ← NUEVO\n'
    '│\n'
    '├── pages/\n'
    '│   ├── login/                       ← Formulario de inicio de sesión\n'
    '│   ├── dashboard/                   ← Tarjetas de módulos según permisos\n'
    '│   ├── presupuestos/\n'
    '│   │   ├── list/  form/  detail/\n'
    '│   ├── pedidos/\n'
    '│   │   ├── list/  detail/\n'
    '│   ├── empresas/\n'
    '│   │   ├── list/  form/  detail/\n'
    '│   ├── contactos/\n'
    '│   │   ├── list/  form/  detail/\n'
    '│   ├── estaciones/\n'
    '│   │   ├── list/  form/  detail/\n'
    '│   ├── usuarios/                    ← NUEVO\n'
    '│   │   ├── list/                    ← Tabla de usuarios (solo admin)\n'
    '│   │   └── form/                    ← Crear/editar usuario + permisos\n'
    '│   ├── contratos/                   ← NUEVO\n'
    '│   │   ├── list/                    ← Tabla con estado vigente/caducado\n'
    '│   │   └── form/                    ← Crear/editar contrato\n'
    '│   ├── tarifario/                   ← NUEVO\n'
    '│   │   ├── list/                    ← Tabla de tarifarios\n'
    '│   │   ├── form/                    ← Crear nuevo tarifario\n'
    '│   │   └── detail/                  ← Ver/editar + tabla servicios inline\n'
    '│   └── facturas/                    ← NUEVO\n'
    '│       ├── list/                    ← Tabla con estado de cobro\n'
    '│       ├── form/                    ← Crear factura + líneas + pedidos\n'
    '│       └── detail/                  ← Ver factura + registrar cobros\n'
    '│\n'
    '└── shared/\n'
    '    └── navbar/                      ← Barra de navegación\n'
    '        ├── navbar.component.ts      ← Muestra módulos según permisos/rol\n'
    '        └── navbar.component.css'
)

para26 = doc.paragraphs[26]
replace_para_text(para26, NEW_FRONTEND)
print('✓ Estructura de carpetas frontend actualizada')

# ─────────────────────────────────────────────────────────
# 5. TABLA 1 — Añadir nuevos endpoints
# ─────────────────────────────────────────────────────────
table1 = doc.tables[1]
NEW_ENDPOINTS = [
    ('GET',    '/usuarios',                                    'Sí (admin)', 'Lista todos los usuarios'),
    ('POST',   '/usuarios',                                    'Sí (admin)', 'Crea un nuevo usuario'),
    ('GET',    '/usuarios/{id}',                               'Sí (admin)', 'Detalle de un usuario'),
    ('PUT',    '/usuarios/{id}',                               'Sí (admin)', 'Actualiza usuario y permisos'),
    ('DELETE', '/usuarios/{id}',                               'Sí (admin)', 'Elimina usuario (no el propio)'),
    ('GET',    '/contratos',                                   'Sí', 'Lista todos los contratos'),
    ('POST',   '/contratos',                                   'Sí', 'Crea un nuevo contrato'),
    ('GET',    '/contratos/{id}',                              'Sí', 'Detalle de un contrato'),
    ('PUT',    '/contratos/{id}',                              'Sí', 'Actualiza un contrato'),
    ('DELETE', '/contratos/{id}',                              'Sí', 'Elimina un contrato'),
    ('GET',    '/tarifario',                                   'Sí', 'Lista todos los tarifarios'),
    ('POST',   '/tarifario',                                   'Sí', 'Crea un nuevo tarifario'),
    ('GET',    '/tarifario/{id}',                              'Sí', 'Detalle con todos sus servicios'),
    ('PUT',    '/tarifario/{id}',                              'Sí', 'Actualiza datos del tarifario'),
    ('DELETE', '/tarifario/{id}',                              'Sí', 'Elimina tarifario y sus servicios'),
    ('GET',    '/tarifario/{id}/servicios',                    'Sí', 'Lista servicios de un tarifario'),
    ('POST',   '/tarifario/{id}/servicios',                    'Sí', 'Añade servicio al tarifario'),
    ('PUT',    '/tarifario/{id}/servicios/{id_servicio}',      'Sí', 'Actualiza un servicio'),
    ('DELETE', '/tarifario/{id}/servicios/{id_servicio}',      'Sí', 'Elimina un servicio'),
    ('GET',    '/facturas',                                    'Sí', 'Lista todas las facturas'),
    ('POST',   '/facturas',                                    'Sí', 'Crea factura con líneas y pedidos'),
    ('GET',    '/facturas/{id}',                               'Sí', 'Detalle: pedidos, líneas y cobros'),
    ('PUT',    '/facturas/{id}',                               'Sí', 'Actualiza datos de la factura'),
    ('DELETE', '/facturas/{id}',                               'Sí', 'Elimina factura y todos sus datos'),
    ('GET',    '/facturas/{id}/cobros',                        'Sí', 'Lista cobros de una factura'),
    ('POST',   '/facturas/{id}/cobros',                        'Sí', 'Registra un cobro'),
    ('DELETE', '/facturas/{id}/cobros/{id_cobro}',             'Sí', 'Elimina un cobro'),
]
for ep in NEW_ENDPOINTS:
    row = table1.add_row()
    for i, txt in enumerate(ep):
        row.cells[i].text = txt
print(f'✓ Tabla de endpoints ampliada con {len(NEW_ENDPOINTS)} nuevas filas')

# ─────────────────────────────────────────────────────────
# 6. TABLA 2 — Añadir nuevas rutas frontend
# ─────────────────────────────────────────────────────────
table2 = doc.tables[2]
NEW_ROUTES = [
    ('/usuarios',              'UsuariosListComponent',       'Sí (solo admin)'),
    ('/usuarios/new',          'UsuariosFormComponent',       'Sí (solo admin)'),
    ('/usuarios/:id/edit',     'UsuariosFormComponent',       'Sí (solo admin)'),
    ('/contratos',             'ContratosListComponent',      'Sí + permiso contratos'),
    ('/contratos/new',         'ContratosFormComponent',      'Sí + permiso contratos'),
    ('/contratos/:id/edit',    'ContratosFormComponent',      'Sí + permiso contratos'),
    ('/tarifario',             'TarifarioListComponent',      'Sí + permiso tarifario'),
    ('/tarifario/new',         'TarifarioFormComponent',      'Sí + permiso tarifario'),
    ('/tarifario/:id',         'TarifarioDetailComponent',    'Sí + permiso tarifario'),
    ('/facturas',              'FacturasListComponent',       'Sí + permiso facturas'),
    ('/facturas/new',          'FacturasFormComponent',       'Sí + permiso facturas'),
    ('/facturas/:id',          'FacturasDetailComponent',     'Sí + permiso facturas'),
]
for rt in NEW_ROUTES:
    row = table2.add_row()
    for i, txt in enumerate(rt):
        row.cells[i].text = txt
print(f'✓ Tabla de rutas ampliada con {len(NEW_ROUTES)} nuevas filas')

# ─────────────────────────────────────────────────────────
# 7. TABLA 3 (users) — Añadir campo permisos
# ─────────────────────────────────────────────────────────
table3 = doc.tables[3]
row = table3.add_row()
row.cells[0].text = 'permisos'
row.cells[1].text = 'TEXT'
row.cells[2].text = 'Sí'
row.cells[3].text = ('Permisos por módulo en formato JSON. '
                     'Ejemplo: {"presupuestos": true, "pedidos": true, '
                     '"facturas": false, "contratos": true, '
                     '"tarifario": true, "usuarios": false}. '
                     'Los admin tienen siempre acceso total.')
print('✓ Tabla users ampliada con campo permisos')

# ─────────────────────────────────────────────────────────
# 8. BD — Insertar nuevas tablas después de la tabla de es
# ─────────────────────────────────────────────────────────
NEW_DB_TABLES = [
    ('contratos', [
        ('id_contrato',     'INTEGER PK AUTOINCREMENT',    'No',  'Identificador único'),
        ('id_empresa',      'INTEGER FK → empresas',       'Sí',  'Empresa asociada al contrato'),
        ('numero_contrato', 'TEXT',                        'Sí',  'Número identificador del contrato'),
        ('nombre',          'TEXT',                        'Sí',  'Nombre descriptivo del contrato'),
        ('id_tarifario',    'INTEGER FK → tarifario',      'Sí',  'Tarifario aplicable'),
        ('fecha_inicio',    'DATE',                        'Sí',  'Fecha de inicio de vigencia'),
        ('fecha_fin',       'DATE',                        'Sí',  'Fecha de fin de vigencia'),
        ('descripcion',     'TEXT',                        'Sí',  'Descripción adicional'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
    ('tarifario', [
        ('id_tarifario',    'INTEGER PK AUTOINCREMENT',    'No',  'Identificador único'),
        ('nombre_tarifario','TEXT',                        'No',  'Nombre del tarifario'),
        ('id_empresa',      'INTEGER',                     'Sí',  'Empresa asociada'),
        ('fecha_tarifario', 'DATE',                        'Sí',  'Fecha de inicio de vigencia'),
        ('fecha_fin',       'DATE',                        'Sí',  'Fecha de fin de vigencia'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
    ('tarifario_servicios', [
        ('id_tarifario',    'INTEGER FK (PK1) → tarifario','No',  'Clave primaria compuesta (parte 1)'),
        ('id_servicio',     'INTEGER (PK2)',                'No',  'Clave primaria compuesta (parte 2, autoincremental por tarifario)'),
        ('codigo_servicio', 'TEXT',                        'Sí',  'Código del servicio'),
        ('numero_tarifa',   'TEXT',                        'Sí',  'Número de tarifa'),
        ('nombre_servicio', 'TEXT',                        'No',  'Nombre del servicio'),
        ('precio_unitario', 'REAL',                        'No',  'Precio unitario'),
        ('precio_anterior', 'REAL',                        'Sí',  'Precio anterior (histórico)'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
    ('facturas', [
        ('id_factura',      'INTEGER PK AUTOINCREMENT',    'No',  'Identificador único'),
        ('numero_factura',  'TEXT',                        'Sí',  'Número de factura'),
        ('id_proyecto',     'INTEGER',                     'Sí',  'Referencia a proyecto'),
        ('id_empresa',      'INTEGER',                     'Sí',  'Empresa facturada'),
        ('id_es',           'INTEGER',                     'Sí',  'Estación de servicio asociada'),
        ('id_contrato',     'INTEGER FK → contratos',      'Sí',  'Contrato asociado'),
        ('id_tarifario',    'INTEGER FK → tarifario',      'Sí',  'Tarifario aplicado'),
        ('id_direccion',    'INTEGER',                     'Sí',  'Dirección de facturación'),
        ('fecha_factura',   'DATE',                        'Sí',  'Fecha de la factura'),
        ('fecha_solicitud', 'DATE',                        'Sí',  'Fecha de solicitud'),
        ('importe_total',   'REAL',                        'Sí',  'Importe total (suma de líneas)'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
    ('factura_pedidos', [
        ('id_factura_pedido','INTEGER PK AUTOINCREMENT',   'No',  'Identificador único'),
        ('id_factura',      'INTEGER FK → facturas',       'No',  'Factura a la que pertenece'),
        ('id_pedido',       'INTEGER FK → pedidos',        'No',  'Pedido vinculado'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
    ('lineas_factura', [
        ('id_linea_factura','INTEGER PK AUTOINCREMENT',    'No',  'Identificador único'),
        ('id_factura',      'INTEGER FK → facturas',       'No',  'Factura a la que pertenece'),
        ('id_servicio',     'INTEGER',                     'Sí',  'Servicio del tarifario'),
        ('id_tarifario',    'INTEGER',                     'Sí',  'Tarifario aplicado'),
        ('unidades',        'REAL',                        'Sí',  'Cantidad de unidades'),
        ('precio_unitario', 'REAL',                        'Sí',  'Precio unitario aplicado'),
        ('importe',         'REAL',                        'Sí',  'Importe = unidades × precio_unitario'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
    ('cobros', [
        ('id_cobro',        'INTEGER PK AUTOINCREMENT',    'No',  'Identificador único'),
        ('id_factura',      'INTEGER FK → facturas',       'No',  'Factura a la que pertenece'),
        ('importe',         'REAL',                        'No',  'Importe cobrado'),
        ('fecha',           'DATE',                        'No',  'Fecha del cobro'),
        ('tipologia_cobro', 'TEXT',                        'Sí',  'Tipo (transferencia, cheque, etc.)'),
        ('cuenta_bancaria', 'TEXT',                        'Sí',  'Cuenta bancaria utilizada'),
        ('created_at',      'DATETIME',                    'Sí',  'Fecha de creación'),
        ('updated_at',      'DATETIME',                    'Sí',  'Fecha de última modificación'),
    ]),
]

# Insertar después de la tabla 11 (es / estaciones)
ref = doc.tables[11]._tbl
for tname, fields in NEW_DB_TABLES:
    h_elem, tbl_elem = make_db_section(tname, fields)
    ref.addnext(tbl_elem)        # tabla justo después de ref
    tbl_elem.addprevious(h_elem) # heading justo antes de la tabla
    ref = tbl_elem               # siguiente inserción después de esta tabla
print(f'✓ {len(NEW_DB_TABLES)} nuevas tablas de BD insertadas')

# ─────────────────────────────────────────────────────────
# 9. RELACIONES — actualizar texto
# ─────────────────────────────────────────────────────────
NEW_RELATIONS = (
    'users\n'
    '  (sin FK salientes)\n\n'
    'presupuestos\n'
    '  └─ presupuestos_lineas  (id_presupuesto → presupuestos.id_presupuesto)\n'
    '  └─ pedidos              (id_presupuesto → presupuestos.id_presupuesto)\n\n'
    'pedidos\n'
    '  └─ lineas_pedido        (id_pedido → pedidos.id_pedido)\n'
    '  └─ factura_pedidos      (id_pedido → pedidos.id_pedido)\n\n'
    'empresas\n'
    '  └─ contactos_empresas   (id_empresa → empresas.id_empresa)\n'
    '  └─ es                   (id_empresa → empresas.id_empresa)\n'
    '  └─ contratos            (id_empresa → empresas.id_empresa)\n\n'
    'contactos\n'
    '  └─ contactos_empresas   (id_contacto → contactos.id_contacto)\n\n'
    'contactos_empresas\n'
    '  ├─ id_empresa  → empresas.id_empresa\n'
    '  └─ id_contacto → contactos.id_contacto\n\n'
    'contratos\n'
    '  ├─ id_empresa   → empresas.id_empresa\n'
    '  └─ id_tarifario → tarifario.id_tarifario\n\n'
    'tarifario\n'
    '  └─ tarifario_servicios  (id_tarifario → tarifario.id_tarifario)\n\n'
    'facturas\n'
    '  ├─ id_contrato  → contratos.id_contrato\n'
    '  ├─ id_tarifario → tarifario.id_tarifario\n'
    '  ├─ factura_pedidos  (id_factura → facturas.id_factura)\n'
    '  ├─ lineas_factura   (id_factura → facturas.id_factura)\n'
    '  └─ cobros           (id_factura → facturas.id_factura)'
)

# Buscar el párrafo de relaciones
for p in doc.paragraphs:
    if p.text.startswith('users') and 'FK' in p.text:
        replace_para_text(p, NEW_RELATIONS)
        print('✓ Diagrama de relaciones actualizado')
        break

# ─────────────────────────────────────────────────────────
# 10. NUEVAS SECCIONES DE MÓDULOS (al final del documento)
# ─────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.clear(); p.add_run(text)

def h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.clear(); p.add_run(text)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def normal(text):
    doc.add_paragraph(text, style='Normal')

def api_table(rows_data):
    t = doc.add_table(rows=1, cols=3, style='LightGrid-Accent1')
    for i, hdr in enumerate(['Método', 'Ruta', 'Descripción']):
        t.rows[0].cells[i].text = hdr
    for r in rows_data:
        row = t.add_row()
        for i, val in enumerate(r):
            row.cells[i].text = val

# ── MÓDULO 9: USUARIOS ──────────────────────────────────
h1('9. Módulo de Usuarios')
normal('El módulo de Usuarios permite a los administradores gestionar las cuentas de acceso al ERP. '
       'Solo los usuarios con rol admin pueden acceder. Permite crear, editar y eliminar cuentas, '
       'asignar roles y controlar el acceso a cada módulo mediante permisos individuales.')

h2('9.1  Acceso')
bullet('Solo accesible para usuarios con rol admin.')
bullet('Desde el Dashboard → tarjeta "Usuarios".')
bullet('Desde el menú de navegación superior → Usuarios (visible solo para admin).')

h2('9.2  Listado de Usuarios')
bullet('Tabla con Nombre, Email/Identificador, Rol y Acciones.')
bullet('Las contraseñas nunca se muestran en ningún momento.')
bullet('Desde la lista puedes Editar o Eliminar cada usuario.')

h2('9.3  Crear un Usuario')
bullet('Pulsa "+ Nuevo Usuario".')
bullet('Campos: Nombre (obligatorio), Email o identificador (acepta cualquier string, no solo emails), Contraseña (obligatoria al crear), Rol (user / admin).')
bullet('Sección Permisos: checkbox por cada módulo (presupuestos, pedidos, facturas, contratos, tarifario, usuarios).')
bullet('Al seleccionar rol admin, todos los permisos se activan automáticamente.')
bullet('Al crear con rol admin: todos los permisos en true. Con rol user: todos en false por defecto.')

h2('9.4  Editar un Usuario')
bullet('Desde el listado, pulsa "Editar".')
bullet('Puedes cambiar nombre, identificador, rol y permisos.')
bullet('El campo contraseña es opcional al editar: si se deja vacío, la contraseña no se modifica.')

h2('9.5  Eliminar un Usuario')
bullet('Desde el listado, pulsa "Eliminar".')
bullet('No es posible eliminar el propio usuario que está conectado.')

h2('9.6  Sistema de Permisos')
normal('Los permisos se almacenan como JSON en el campo permisos de la tabla users:')
normal('{"presupuestos": true, "pedidos": true, "facturas": false, "contratos": true, "tarifario": true, "usuarios": false}')
bullet('Módulos sin permiso: el usuario es redirigido al Dashboard si intenta acceder directamente.')
bullet('Módulos con permiso: visibles en la barra de navegación y en el Dashboard.')
bullet('El rol admin siempre tiene acceso completo, independientemente del campo permisos.')
bullet('El guard permisos.guard.ts intercepta las rutas protegidas y verifica el permiso correspondiente.')

h2('9.7  Endpoints API')
api_table([
    ('GET',    '/usuarios',           'Lista todos los usuarios (solo admin)'),
    ('POST',   '/usuarios',           'Crea un nuevo usuario (solo admin)'),
    ('GET',    '/usuarios/{id}',      'Detalle de un usuario (solo admin)'),
    ('PUT',    '/usuarios/{id}',      'Actualiza usuario y permisos (solo admin)'),
    ('DELETE', '/usuarios/{id}',      'Elimina usuario, excepto el propio (solo admin)'),
])

# ── MÓDULO 10: CONTRATOS ─────────────────────────────────
h1('10. Módulo de Contratos')
normal('El módulo de Contratos gestiona los acuerdos entre la empresa y sus clientes. '
       'Cada contrato se vincula a una empresa y opcionalmente a un tarifario. '
       'El sistema calcula automáticamente si el contrato está vigente o caducado según la fecha de fin.')

h2('10.1  Acceso')
bullet('Desde el Dashboard → tarjeta "Contratos" (visible si tienes permiso).')
bullet('Desde el menú de navegación superior → Contratos.')
bullet('Requiere permiso de módulo "contratos" o rol admin.')

h2('10.2  Listado de Contratos')
bullet('Tabla con Nº Contrato, Nombre, Empresa, Tarifario, Fecha inicio, Fecha fin y Estado.')
bullet('Estado calculado: "Vigente" si fecha_fin ≥ fecha actual; "Caducado" en caso contrario.')
bullet('Desde la lista puedes Editar o Eliminar cada contrato.')

h2('10.3  Crear un Contrato')
bullet('Pulsa "+ Nuevo Contrato".')
bullet('Campos: Número de contrato, Nombre, Empresa (selector), Tarifario (selector), Fecha de inicio, Fecha de fin, Descripción.')

h2('10.4  Editar un Contrato')
bullet('Pulsa "Editar" desde el listado. Modifica y pulsa "Actualizar".')

h2('10.5  Eliminar un Contrato')
bullet('Pulsa "Eliminar" y confirma en el diálogo emergente.')

h2('10.6  Endpoints API')
api_table([
    ('GET',    '/contratos',       'Lista todos los contratos'),
    ('POST',   '/contratos',       'Crea un nuevo contrato'),
    ('GET',    '/contratos/{id}',  'Detalle de un contrato'),
    ('PUT',    '/contratos/{id}',  'Actualiza un contrato'),
    ('DELETE', '/contratos/{id}',  'Elimina un contrato'),
])

# ── MÓDULO 11: TARIFARIO ─────────────────────────────────
h1('11. Módulo de Tarifario')
normal('El módulo de Tarifario gestiona las listas de precios de los servicios. '
       'Cada tarifario pertenece a una empresa y contiene servicios con código, nombre y precio unitario. '
       'Desde la pantalla de detalle se gestionan los servicios directamente sin navegar a otra página.')

h2('11.1  Acceso')
bullet('Desde el Dashboard → tarjeta "Tarifario" (visible si tienes permiso).')
bullet('Desde el menú de navegación superior → Tarifario.')
bullet('Requiere permiso de módulo "tarifario" o rol admin.')

h2('11.2  Listado de Tarifarios')
bullet('Tabla con Nombre, Empresa, Fecha de vigencia, Fecha de fin y Acciones.')
bullet('Desde la lista puedes Ver/gestionar servicios o Eliminar.')

h2('11.3  Crear un Tarifario')
bullet('Pulsa "+ Nuevo Tarifario".')
bullet('Campos: Nombre (obligatorio), Empresa (selector), Fecha de vigencia, Fecha de fin.')
bullet('Al guardar, se redirige automáticamente al detalle para añadir servicios.')

h2('11.4  Gestionar Servicios (Pantalla de Detalle)')
normal('La pantalla de detalle muestra los datos del tarifario y una tabla editable con sus servicios.')
bullet('Editar datos del tarifario: pulsa "Editar datos" y modifica los campos en línea.')
bullet('Añadir servicio: pulsa "+ Añadir servicio", rellena los campos y confirma.')
bullet('Editar servicio: pulsa "Editar" en la fila, modifica directamente en la tabla y pulsa "Guardar".')
bullet('Eliminar servicio: pulsa "Eliminar" en la fila correspondiente.')

h2('11.5  Campos de un Servicio')
bullet('Código: código de identificación del servicio (opcional).')
bullet('Nº tarifa: número de tarifa de referencia (opcional).')
bullet('Nombre del servicio: descripción del servicio (obligatorio).')
bullet('Precio unitario: precio por unidad en euros (obligatorio).')

h2('11.6  Eliminar un Tarifario')
bullet('Desde el listado, pulsa "Eliminar".')
bullet('Se eliminan también todos los servicios asociados (operación cascade).')

h2('11.7  Endpoints API')
api_table([
    ('GET',    '/tarifario',                                  'Lista todos los tarifarios'),
    ('POST',   '/tarifario',                                  'Crea un nuevo tarifario'),
    ('GET',    '/tarifario/{id}',                             'Detalle con todos sus servicios'),
    ('PUT',    '/tarifario/{id}',                             'Actualiza datos del tarifario'),
    ('DELETE', '/tarifario/{id}',                             'Elimina tarifario y sus servicios'),
    ('GET',    '/tarifario/{id}/servicios',                   'Lista servicios de un tarifario'),
    ('POST',   '/tarifario/{id}/servicios',                   'Añade servicio al tarifario'),
    ('PUT',    '/tarifario/{id}/servicios/{id_servicio}',     'Actualiza un servicio'),
    ('DELETE', '/tarifario/{id}/servicios/{id_servicio}',     'Elimina un servicio'),
])

# ── MÓDULO 12: FACTURAS ──────────────────────────────────
h1('12. Módulo de Facturas')
normal('El módulo de Facturas gestiona la facturación del ERP. Permite crear facturas vinculadas a '
       'empresas, contratos y tarifarios, asociar pedidos, definir líneas de factura con precios de '
       'tarifario y registrar cobros. El estado (Pendiente / Cobro parcial / Cobrada) se calcula '
       'automáticamente comparando el importe total con la suma de cobros registrados.')

h2('12.1  Acceso')
bullet('Desde el Dashboard → tarjeta "Facturas" (visible si tienes permiso).')
bullet('Desde el menú de navegación superior → Facturas.')
bullet('Requiere permiso de módulo "facturas" o rol admin.')

h2('12.2  Listado de Facturas')
bullet('Tabla con Nº Factura, Empresa, Fecha, Importe total y Estado de cobro.')
bullet('"Pendiente": sin cobros. "Cobro parcial": cobros < importe total. "Cobrada": cobros ≥ importe total.')
bullet('Desde la lista puedes Ver el detalle o Eliminar cada factura.')

h2('12.3  Crear una Factura')
bullet('Pulsa "+ Nueva Factura".')
bullet('Sección Datos generales: número de factura, empresa (selector), contrato (se filtra automáticamente por empresa), tarifario, fechas de factura y solicitud.')
bullet('Sección Pedidos vinculados: selecciona pedidos existentes para asociar a la factura.')
bullet('Sección Líneas de factura: añade líneas seleccionando tarifario y servicio. El precio unitario se carga automáticamente al elegir el servicio. El importe se calcula al introducir las unidades.')
bullet('El importe total se calcula como la suma de todos los importes de línea.')

h2('12.4  Ver Detalle de una Factura')
bullet('Pulsa "Ver" desde el listado.')
bullet('Muestra datos generales, pedidos vinculados, tabla de líneas con importes y total.')
bullet('Sección Cobros: lista los cobros registrados y permite añadir nuevos con formulario inline.')

h2('12.5  Registrar un Cobro')
bullet('Desde el detalle, pulsa "+ Registrar cobro".')
bullet('Campos: Importe (obligatorio), Fecha (obligatoria), Tipología (transferencia, cheque...), Cuenta bancaria.')
bullet('El estado de cobro se recalcula automáticamente en la vista.')

h2('12.6  Eliminar una Factura')
bullet('Desde el listado, pulsa "Eliminar".')
bullet('Se eliminan también todas las líneas, vínculos con pedidos y cobros (cascade).')

h2('12.7  Endpoints API')
api_table([
    ('GET',    '/facturas',                          'Lista todas las facturas'),
    ('POST',   '/facturas',                          'Crea factura con líneas y pedidos vinculados'),
    ('GET',    '/facturas/{id}',                     'Detalle: pedidos, líneas y cobros'),
    ('PUT',    '/facturas/{id}',                     'Actualiza datos generales de la factura'),
    ('DELETE', '/facturas/{id}',                     'Elimina factura y todos sus datos'),
    ('GET',    '/facturas/{id}/cobros',              'Lista los cobros de una factura'),
    ('POST',   '/facturas/{id}/cobros',              'Registra un nuevo cobro'),
    ('DELETE', '/facturas/{id}/cobros/{id_cobro}',   'Elimina un cobro'),
])

# ─────────────────────────────────────────────────────────
# GUARDAR
# ─────────────────────────────────────────────────────────
doc.save(PATH)
print()
print('✓ Documento guardado en:', PATH)
