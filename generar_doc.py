# Script para generar la documentacion del ERP Ciete en formato Word (.docx)
# Ejecutar desde: c:\Users\Bernardo\Documents\erp-ciete\erp-ciete\erp-ciete\
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F1F5F9')
    pPr = p._p.get_or_add_pPr()
    pPr.append(shading)
    return p

def page_break():
    doc.add_page_break()

def make_api_table(endpoints, color='2563EB'):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, t in enumerate(['Método', 'Ruta', 'Descripción']):
        hdr[i].text = t
        shade_cell(hdr[i], color)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for m, r, d in endpoints:
        row = table.add_row().cells
        row[0].text = m
        row[1].text = r
        row[2].text = d

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ERP CIETE INGENIEROS')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación Técnica y de Usuario')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x37, 0x51, 0x8E)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Versión 3.0  ·  {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ÍNDICE (manual)
# ═══════════════════════════════════════════════════════════════════════════════
h1('Índice')
indices = [
    ('1.',  'Arquitectura y estructura del proyecto'),
    ('2.',  'Esquema de la base de datos'),
    ('3.',  'Guía de instalación'),
    ('4.',  'Credenciales de acceso'),
    ('5.',  'Guía de uso diario'),
    ('6.',  'Módulo de Empresas'),
    ('7.',  'Módulo de Contactos'),
    ('8.',  'Módulo de Estaciones de Servicio'),
    ('9.',  'Módulo de Presupuestos'),
    ('10.', 'Módulo de Pedidos'),
    ('11.', 'Módulo de Contratos'),
    ('12.', 'Módulo de Tarifario'),
    ('13.', 'Módulo de Facturas'),
    ('14.', 'Gestión de Usuarios'),
    ('15.', 'Sistema de permisos y roles'),
]
for num, title in indices:
    p = doc.add_paragraph()
    run = p.add_run(f'  {num}  {title}')
    run.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. ARQUITECTURA Y ESTRUCTURA
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Arquitectura y Estructura del Proyecto')

body(
    'El ERP Ciete es una aplicación web de dos capas: un backend REST API '
    'desarrollado en Python/FastAPI y un frontend SPA (Single Page Application) '
    'en Angular 17. Ambas piezas son completamente independientes y se comunican '
    'mediante HTTP/JSON.'
)

h2('1.1  Stack Tecnológico')
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Capa'
hdr[1].text = 'Tecnología'
hdr[2].text = 'Versión'
for c in hdr:
    shade_cell(c, '1E3A8A')
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c.paragraphs[0].runs[0].font.bold = True

rows_data = [
    ('Backend', 'Python', '3.13'),
    ('Backend', 'FastAPI', '0.135+'),
    ('Backend', 'SQLAlchemy', '2.0+'),
    ('Backend', 'SQLite', '3.x (archivo)'),
    ('Backend', 'bcrypt', '4.x+'),
    ('Backend', 'python-jose', '3.5+'),
    ('Frontend', 'Angular', '17+'),
    ('Frontend', 'TypeScript', '5.4+'),
    ('Frontend', 'Node.js', '18+ / 20+'),
]
for capa, tech, ver in rows_data:
    row = table.add_row().cells
    row[0].text = capa
    row[1].text = tech
    row[2].text = ver

doc.add_paragraph()

h2('1.2  Estructura de Carpetas — Backend')
code_block(
    'backend/\n'
    '├── main.py                  ← Punto de entrada FastAPI. Configura CORS,\n'
    '│                              registra routers y crea el usuario admin\n'
    '│                              al arrancar (startup event).\n'
    '├── database.py              ← Conexión SQLite con SQLAlchemy.\n'
    '│                              Expone get_db() como dependencia.\n'
    '├── auth.py                  ← JWT (python-jose) + hashing bcrypt.\n'
    '│                              Funciones: create_access_token,\n'
    '│                              get_current_user, verify_password,\n'
    '│                              get_password_hash.\n'
    '├── requirements.txt         ← Dependencias Python del proyecto.\n'
    '├── database.sqlite          ← Archivo de base de datos (auto-generado).\n'
    '│\n'
    '├── models/                  ← Modelos SQLAlchemy (tablas de la BD)\n'
    '│   ├── user.py              ← Tabla users (con campo permisos)\n'
    '│   ├── presupuesto.py       ← Tablas presupuestos + presupuestos_lineas\n'
    '│   ├── pedido.py            ← Tablas pedidos + lineas_pedido\n'
    '│   ├── empresa.py           ← Tablas empresas + contactos_empresas\n'
    '│   ├── contacto.py          ← Tabla contactos\n'
    '│   ├── estacion.py          ← Tabla es (estaciones de servicio)\n'
    '│   ├── contrato.py          ← Tabla contratos\n'
    '│   ├── tarifario.py         ← Tablas tarifario + tarifario_servicios\n'
    '│   └── factura.py           ← Tablas facturas + factura_pedidos\n'
    '│                               + lineas_factura + cobros\n'
    '│\n'
    '├── schemas/                 ← Schemas Pydantic (validación + serialización)\n'
    '│   ├── user.py              ← UserOut, UserCreate, Token, LoginRequest\n'
    '│   ├── usuario.py           ← UsuarioCreate, UsuarioUpdate, UsuarioOut\n'
    '│   ├── presupuesto.py       ← PresupuestoOut, PresupuestoCreate, etc.\n'
    '│   ├── pedido.py            ← PedidoOut, LineaPedidoOut\n'
    '│   ├── empresa.py           ← EmpresaOut, EmpresaCreate,\n'
    '│   │                          ContactoEmpresaOut, ContactoEmpresaCreate\n'
    '│   ├── contacto.py          ← ContactoOut, ContactoCreate\n'
    '│   ├── estacion.py          ← EstacionOut, EstacionCreate\n'
    '│   ├── contrato.py          ← ContratoOut, ContratoCreate\n'
    '│   ├── tarifario.py         ← TarifarioOut, TarifarioCreate,\n'
    '│   │                          TarifarioServicioOut, TarifarioServicioCreate\n'
    '│   └── factura.py           ← FacturaOut, FacturaCreate,\n'
    '│                               CobroOut, CobroCreate\n'
    '│\n'
    '└── routers/                 ← Endpoints REST organizados por área\n'
    '    ├── auth.py              ← POST /auth/login, /auth/registro, GET /auth/me\n'
    '    ├── presupuestos.py      ← CRUD presupuestos + cambio estado + convertir\n'
    '    ├── pedidos.py           ← GET /pedidos, GET /pedidos/{id}\n'
    '    ├── empresas.py          ← CRUD empresas + gestión contactos-empresa\n'
    '    ├── contactos.py         ← CRUD contactos\n'
    '    ├── estaciones.py        ← CRUD estaciones de servicio\n'
    '    ├── contratos.py         ← CRUD contratos\n'
    '    ├── tarifario.py         ← CRUD tarifarios + CRUD servicios por tarifario\n'
    '    ├── facturas.py          ← CRUD facturas + pedidos vinculados + cobros\n'
    '    └── usuarios.py          ← CRUD usuarios (solo admin) + gestión permisos'
)

doc.add_paragraph()
h2('1.3  Endpoints de la API')
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, t in enumerate(['Método', 'Ruta', 'Auth', 'Descripción']):
    hdr[i].text = t
    shade_cell(hdr[i], '2563EB')
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr[i].paragraphs[0].runs[0].font.bold = True

endpoints = [
    # Auth
    ('POST', '/auth/login',                              'No',        'Login → devuelve JWT'),
    ('POST', '/auth/registro',                           'No',        'Registra nuevo usuario'),
    ('GET',  '/auth/me',                                 'Sí',        'Datos del usuario actual'),
    # Presupuestos
    ('GET',  '/presupuestos',                            'Sí',        'Lista todos los presupuestos'),
    ('POST', '/presupuestos',                            'Sí',        'Crea presupuesto con líneas'),
    ('GET',  '/presupuestos/{id}',                       'Sí',        'Detalle de un presupuesto con líneas'),
    ('PATCH','/presupuestos/{id}/estado',                'Sí',        'Cambia estado del presupuesto'),
    ('POST', '/presupuestos/{id}/convertir',             'Sí',        'Convierte presupuesto a pedido'),
    # Pedidos
    ('GET',  '/pedidos',                                 'Sí',        'Lista todos los pedidos'),
    ('GET',  '/pedidos/{id}',                            'Sí',        'Detalle de un pedido'),
    # Empresas
    ('GET',  '/empresas',                                'Sí',        'Lista todas las empresas'),
    ('POST', '/empresas',                                'Sí',        'Crea una nueva empresa'),
    ('GET',  '/empresas/{id}',                           'Sí',        'Detalle de una empresa'),
    ('PUT',  '/empresas/{id}',                           'Sí',        'Actualiza una empresa'),
    ('DELETE','/empresas/{id}',                          'Sí',        'Elimina una empresa'),
    ('GET',  '/empresas/{id}/contactos',                 'Sí',        'Contactos de la empresa'),
    ('POST', '/empresas/{id}/contactos',                 'Sí',        'Vincula contacto a empresa'),
    # Contactos
    ('GET',  '/contactos',                               'Sí',        'Lista todos los contactos'),
    ('POST', '/contactos',                               'Sí',        'Crea un nuevo contacto'),
    ('GET',  '/contactos/{id}',                          'Sí',        'Detalle de un contacto'),
    ('PUT',  '/contactos/{id}',                          'Sí',        'Actualiza un contacto'),
    ('DELETE','/contactos/{id}',                         'Sí',        'Elimina un contacto'),
    # Estaciones
    ('GET',  '/estaciones',                              'Sí',        'Lista todas las estaciones'),
    ('POST', '/estaciones',                              'Sí',        'Crea una nueva estación'),
    ('GET',  '/estaciones/{id}',                         'Sí',        'Detalle de una estación'),
    ('PUT',  '/estaciones/{id}',                         'Sí',        'Actualiza una estación'),
    ('DELETE','/estaciones/{id}',                        'Sí',        'Elimina una estación'),
    # Contratos
    ('GET',  '/contratos',                               'Sí',        'Lista todos los contratos'),
    ('POST', '/contratos',                               'Sí',        'Crea un nuevo contrato'),
    ('GET',  '/contratos/{id}',                          'Sí',        'Detalle de un contrato'),
    ('PUT',  '/contratos/{id}',                          'Sí',        'Actualiza un contrato'),
    ('DELETE','/contratos/{id}',                         'Sí',        'Elimina un contrato'),
    # Tarifario
    ('GET',  '/tarifario',                               'Sí',        'Lista todos los tarifarios'),
    ('POST', '/tarifario',                               'Sí',        'Crea un nuevo tarifario'),
    ('GET',  '/tarifario/{id}',                          'Sí',        'Detalle con sus servicios'),
    ('PUT',  '/tarifario/{id}',                          'Sí',        'Actualiza un tarifario'),
    ('DELETE','/tarifario/{id}',                         'Sí',        'Elimina un tarifario'),
    ('GET',  '/tarifario/{id}/servicios',                'Sí',        'Lista servicios del tarifario'),
    ('POST', '/tarifario/{id}/servicios',                'Sí',        'Añade un servicio al tarifario'),
    ('PUT',  '/tarifario/{id_t}/servicios/{id_s}',       'Sí',        'Actualiza un servicio'),
    ('DELETE','/tarifario/{id_t}/servicios/{id_s}',      'Sí',        'Elimina un servicio'),
    # Facturas
    ('GET',  '/facturas',                                'Sí',        'Lista todas las facturas'),
    ('POST', '/facturas',                                'Sí',        'Crea factura con líneas y pedidos'),
    ('GET',  '/facturas/{id}',                           'Sí',        'Detalle completo con cobros'),
    ('PUT',  '/facturas/{id}',                           'Sí',        'Actualiza datos generales'),
    ('DELETE','/facturas/{id}',                          'Sí',        'Elimina factura y sus líneas'),
    ('GET',  '/facturas/{id}/cobros',                    'Sí',        'Lista cobros de la factura'),
    ('POST', '/facturas/{id}/cobros',                    'Sí',        'Registra un cobro'),
    ('DELETE','/facturas/{id}/cobros/{id_cobro}',        'Sí',        'Elimina un cobro'),
    # Usuarios
    ('GET',  '/usuarios',                                'Admin',     'Lista todos los usuarios'),
    ('POST', '/usuarios',                                'Admin',     'Crea un nuevo usuario'),
    ('GET',  '/usuarios/{id}',                           'Admin',     'Detalle de un usuario'),
    ('PUT',  '/usuarios/{id}',                           'Admin',     'Edita nombre, email, rol, permisos'),
    ('DELETE','/usuarios/{id}',                          'Admin',     'Elimina usuario (no el propio)'),
]
for metodo, ruta, auth, desc in endpoints:
    row = table.add_row().cells
    row[0].text = metodo
    row[1].text = ruta
    row[2].text = auth
    row[3].text = desc

doc.add_paragraph()
h2('1.4  Estructura de Carpetas — Frontend')
code_block(
    'frontend/src/app/\n'
    '├── app.component.ts         ← Componente raíz (solo renderiza <router-outlet>)\n'
    '├── app.config.ts            ← Configura provideHttpClient + provideRouter\n'
    '├── app.routes.ts            ← Definición de rutas con lazy loading\n'
    '│\n'
    '├── core/\n'
    '│   ├── services/\n'
    '│   │   ├── auth.service.ts          ← login, logout, getToken, getMe\n'
    '│   │   ├── presupuestos.service.ts  ← getAll, getOne, create,\n'
    '│   │   │                               cambiarEstado, convertir\n'
    '│   │   ├── pedidos.service.ts       ← getAll, getOne\n'
    '│   │   ├── empresas.service.ts      ← getAll, getOne, create,\n'
    '│   │   │                               update, delete\n'
    '│   │   ├── contactos.service.ts     ← getAll, getOne, create,\n'
    '│   │   │                               update, delete\n'
    '│   │   ├── estaciones.service.ts    ← getAll, getOne, create,\n'
    '│   │   │                               update, delete\n'
    '│   │   ├── contratos.service.ts     ← getAll, getOne, create,\n'
    '│   │   │                               update, delete\n'
    '│   │   ├── tarifario.service.ts     ← getAll, getOne, create, update,\n'
    '│   │   │                               delete, getServicios,\n'
    '│   │   │                               addServicio, updateServicio,\n'
    '│   │   │                               deleteServicio\n'
    '│   │   ├── facturas.service.ts      ← getAll, getOne, create, update,\n'
    '│   │   │                               delete, getCobros, addCobro,\n'
    '│   │   │                               deleteCobro\n'
    '│   │   └── usuarios.service.ts      ← getAll, getOne, create,\n'
    '│   │                                   update, delete\n'
    '│   ├── interceptors/\n'
    '│   │   └── auth.interceptor.ts      ← Añade Bearer token a cada petición\n'
    '│   └── guards/\n'
    '│       ├── auth.guard.ts            ← Redirige a /login si no hay token\n'
    '│       └── permisos.guard.ts        ← Comprueba permiso del módulo;\n'
    '│                                       redirige al dashboard si denegado\n'
    '│\n'
    '├── pages/\n'
    '│   ├── login/                       ← Formulario de inicio de sesión\n'
    '│   ├── dashboard/                   ← Pantalla principal con accesos\n'
    '│   │                                   filtrados por permisos del usuario\n'
    '│   ├── presupuestos/\n'
    '│   │   ├── list/                    ← Tabla de presupuestos\n'
    '│   │   ├── form/                    ← Nuevo presupuesto con líneas\n'
    '│   │   └── detail/                  ← Detalle + estado + convertir\n'
    '│   ├── pedidos/\n'
    '│   │   ├── list/                    ← Tabla de pedidos\n'
    '│   │   └── detail/                  ← Detalle completo del pedido\n'
    '│   ├── empresas/\n'
    '│   │   ├── list/                    ← Tabla de empresas\n'
    '│   │   ├── form/                    ← Crear / editar empresa\n'
    '│   │   └── detail/                  ← Detalle de la empresa\n'
    '│   ├── contactos/\n'
    '│   │   ├── list/                    ← Tabla de contactos\n'
    '│   │   ├── form/                    ← Crear / editar contacto\n'
    '│   │   └── detail/                  ← Detalle del contacto\n'
    '│   ├── estaciones/\n'
    '│   │   ├── list/                    ← Tabla de estaciones\n'
    '│   │   ├── form/                    ← Crear / editar estación\n'
    '│   │   └── detail/                  ← Detalle de la estación\n'
    '│   ├── contratos/\n'
    '│   │   ├── list/                    ← Tabla: nº contrato, empresa,\n'
    '│   │   │                               tarifario, fechas, estado\n'
    '│   │   └── form/                    ← Crear / editar contrato\n'
    '│   ├── tarifario/\n'
    '│   │   ├── list/                    ← Tabla de tarifarios\n'
    '│   │   └── detail/                  ← Ficha + tabla editable de servicios\n'
    '│   ├── facturas/\n'
    '│   │   ├── list/                    ← Tabla con estado de cobro calculado\n'
    '│   │   ├── form/                    ← Crear factura (empresa, proyecto,\n'
    '│   │   │                               contrato, líneas, pedidos)\n'
    '│   │   └── detail/                  ← Detalle + cobros inline\n'
    '│   └── usuarios/\n'
    '│       ├── list/                    ← Tabla: nombre, email/usuario, rol\n'
    '│       └── form/                    ← Crear / editar usuario + permisos\n'
    '│\n'
    '└── shared/\n'
    '    └── navbar/                      ← Barra de navegación compartida\n'
    '                                        (muestra Usuarios solo si role=admin)'
)

doc.add_paragraph()
h2('1.5  Rutas del Frontend')
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, t in enumerate(['Ruta', 'Componente', 'Guards', 'Módulo permiso']):
    hdr[i].text = t
    shade_cell(hdr[i], '2563EB')
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr[i].paragraphs[0].runs[0].font.bold = True

rutas = [
    ('/login',                 'LoginComponent',              'Ninguno',             '—'),
    ('/dashboard',             'DashboardComponent',          'AuthGuard',           '—'),
    ('/presupuestos',          'PresupuestosListComponent',   'AuthGuard',           '—'),
    ('/presupuestos/new',      'PresupuestosFormComponent',   'AuthGuard',           '—'),
    ('/presupuestos/:id',      'PresupuestosDetailComponent', 'AuthGuard',           '—'),
    ('/pedidos',               'PedidosListComponent',        'AuthGuard',           '—'),
    ('/pedidos/:id',           'PedidosDetailComponent',      'AuthGuard',           '—'),
    ('/empresas',              'EmpresasListComponent',       'AuthGuard',           '—'),
    ('/empresas/new',          'EmpresasFormComponent',       'AuthGuard',           '—'),
    ('/empresas/:id',          'EmpresasDetailComponent',     'AuthGuard',           '—'),
    ('/empresas/:id/edit',     'EmpresasFormComponent',       'AuthGuard',           '—'),
    ('/contactos',             'ContactosListComponent',      'AuthGuard',           '—'),
    ('/contactos/new',         'ContactosFormComponent',      'AuthGuard',           '—'),
    ('/contactos/:id',         'ContactosDetailComponent',    'AuthGuard',           '—'),
    ('/contactos/:id/edit',    'ContactosFormComponent',      'AuthGuard',           '—'),
    ('/estaciones',            'EstacionesListComponent',     'AuthGuard',           '—'),
    ('/estaciones/new',        'EstacionesFormComponent',     'AuthGuard',           '—'),
    ('/estaciones/:id',        'EstacionesDetailComponent',   'AuthGuard',           '—'),
    ('/estaciones/:id/edit',   'EstacionesFormComponent',     'AuthGuard',           '—'),
    ('/contratos',             'ContratosListComponent',      'AuthGuard+Permisos',  'contratos'),
    ('/contratos/new',         'ContratosFormComponent',      'AuthGuard+Permisos',  'contratos'),
    ('/contratos/:id/edit',    'ContratosFormComponent',      'AuthGuard+Permisos',  'contratos'),
    ('/tarifario',             'TarifarioListComponent',      'AuthGuard+Permisos',  'tarifario'),
    ('/tarifario/new',         'TarifarioFormComponent',      'AuthGuard+Permisos',  'tarifario'),
    ('/tarifario/:id',         'TarifarioDetailComponent',    'AuthGuard+Permisos',  'tarifario'),
    ('/facturas',              'FacturasListComponent',       'AuthGuard+Permisos',  'facturas'),
    ('/facturas/new',          'FacturasFormComponent',       'AuthGuard+Permisos',  'facturas'),
    ('/facturas/:id',          'FacturasDetailComponent',     'AuthGuard+Permisos',  'facturas'),
    ('/usuarios',              'UsuariosListComponent',       'AuthGuard',           'admin'),
    ('/usuarios/new',          'UsuariosFormComponent',       'AuthGuard',           'admin'),
    ('/usuarios/:id/edit',     'UsuariosFormComponent',       'AuthGuard',           'admin'),
]
for r, c, g, m in rutas:
    row = table.add_row().cells
    row[0].text = r
    row[1].text = c
    row[2].text = g
    row[3].text = m

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ESQUEMA DE LA BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Esquema de la Base de Datos')
body(
    'La base de datos es un único archivo SQLite ubicado en backend/database.sqlite. '
    'Se crea automáticamente al arrancar el servidor por primera vez. '
    'Contiene 15 tablas:'
)

def db_table(title, fields, color='1E3A8A'):
    h2(f'Tabla: {title}')
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Campo', 'Tipo', 'Nulo', 'Descripción']):
        t.rows[0].cells[i].text = h
        shade_cell(t.rows[0].cells[i], color)
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for row in fields:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

# users
db_table('users', [
    ('id',                 'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('name',               'VARCHAR',                  'No',  'Nombre completo'),
    ('email',              'VARCHAR UNIQUE',            'No',  'Email o identificador de login'),
    ('email_verified_at',  'DATETIME',                 'Sí',  'Verificación de email'),
    ('role',               'VARCHAR',                  'No',  "'user' o 'admin'"),
    ('password',           'VARCHAR',                  'No',  'Hash bcrypt de la contraseña'),
    ('permisos',           'TEXT',                     'Sí',  'JSON con permisos por módulo'),
    ('remember_token',     'VARCHAR',                  'Sí',  'Token sesión persistente'),
    ('created_at',         'DATETIME',                 'Sí',  'Fecha de creación'),
    ('updated_at',         'DATETIME',                 'Sí',  'Última modificación'),
])

# presupuestos
db_table('presupuestos', [
    ('id_presupuesto',        'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('id_proyecto',           'INTEGER',                  'Sí',  'Referencia a proyecto externo'),
    ('id_empresa',            'INTEGER',                  'Sí',  'Referencia a empresa'),
    ('id_contactos_empresas', 'INTEGER',                  'Sí',  'Referencia a contacto'),
    ('id_es',                 'INTEGER',                  'Sí',  'Referencia a ES'),
    ('id_tarifario',          'INTEGER',                  'Sí',  'Referencia a tarifario'),
    ('fecha_presupuesto',     'DATE',                     'Sí',  'Fecha del presupuesto'),
    ('estado',                'VARCHAR',                  'No',  'borrador/enviado/aprobado/rechazado'),
    ('created_at',            'DATETIME',                 'Sí',  'Fecha de creación'),
    ('updated_at',            'DATETIME',                 'Sí',  'Última modificación'),
])

# presupuestos_lineas
db_table('presupuestos_lineas', [
    ('id_linea_presupuesto', 'INTEGER PK AUTOINCREMENT',  'No', 'Identificador único'),
    ('id_presupuesto',       'INTEGER FK → presupuestos', 'No', 'Presupuesto al que pertenece'),
    ('id_tarifario',         'INTEGER',                   'Sí', 'Referencia a tarifario'),
    ('id_servicio',          'INTEGER',                   'Sí', 'Referencia al servicio'),
    ('unidades',             'NUMERIC',                   'Sí', 'Cantidad de unidades'),
    ('created_at',           'DATETIME',                  'Sí', 'Fecha de creación'),
    ('updated_at',           'DATETIME',                  'Sí', 'Última modificación'),
])

# pedidos
db_table('pedidos', [
    ('id_pedido',                  'INTEGER PK AUTOINCREMENT',  'No', 'Identificador único'),
    ('id_presupuesto',             'INTEGER FK → presupuestos', 'Sí', 'Presupuesto de origen'),
    ('id_proyecto',                'INTEGER',                   'Sí', 'Referencia a proyecto'),
    ('id_empresa',                 'INTEGER',                   'Sí', 'Referencia a empresa'),
    ('id_es',                      'INTEGER',                   'Sí', 'Referencia a ES'),
    ('id_tarifario',               'INTEGER',                   'Sí', 'Referencia a tarifario'),
    ('fecha_solicitud_pedido',     'DATE',                      'Sí', 'Fecha de solicitud del pedido'),
    ('fecha_solicitud_autofactura','DATE',                      'Sí', 'Fecha de solicitud de autofactura'),
    ('fecha_recepcion_pedido',     'DATE',                      'Sí', 'Fecha de recepción'),
    ('estado',                     'VARCHAR',                   'No', 'pendiente/completado/etc.'),
    ('created_at',                 'DATETIME',                  'Sí', 'Fecha de creación'),
    ('updated_at',                 'DATETIME',                  'Sí', 'Última modificación'),
])

# lineas_pedido
db_table('lineas_pedido', [
    ('id_linea_pedido', 'INTEGER PK AUTOINCREMENT', 'No', 'Identificador único'),
    ('id_pedido',       'INTEGER FK → pedidos',     'No', 'Pedido al que pertenece'),
    ('id_servicio',     'INTEGER',                  'Sí', 'Referencia al servicio'),
    ('unidades',        'NUMERIC',                  'Sí', 'Cantidad de unidades'),
    ('created_at',      'DATETIME',                 'Sí', 'Fecha de creación'),
    ('updated_at',      'DATETIME',                 'Sí', 'Última modificación'),
])

# empresas
db_table('empresas', [
    ('id_empresa',   'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('nombre',       'TEXT',                      'No',  'Nombre de la empresa'),
    ('razon_social', 'TEXT',                      'Sí',  'Razón social'),
    ('cif',          'TEXT UNIQUE',               'Sí',  'CIF de la empresa'),
    ('id_direccion', 'INTEGER',                   'Sí',  'Referencia a dirección'),
    ('id_telefono',  'INTEGER',                   'Sí',  'Referencia a teléfono'),
    ('id_email',     'INTEGER',                   'Sí',  'Referencia a email'),
    ('created_at',   'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',   'DATETIME',                  'Sí',  'Última modificación'),
])

# contactos
db_table('contactos', [
    ('id_contacto',  'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('nombre',       'TEXT',                      'No',  'Nombre del contacto'),
    ('apellido1',    'TEXT',                      'Sí',  'Primer apellido'),
    ('apellido2',    'TEXT',                      'Sí',  'Segundo apellido'),
    ('id_direccion', 'INTEGER',                   'Sí',  'Referencia a dirección'),
    ('id_telefono',  'INTEGER',                   'Sí',  'Referencia a teléfono'),
    ('id_email',     'INTEGER',                   'Sí',  'Referencia a email'),
    ('created_at',   'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',   'DATETIME',                  'Sí',  'Última modificación'),
])

# contactos_empresas
db_table('contactos_empresas', [
    ('id_contactos_empresas', 'INTEGER PK AUTOINCREMENT',  'No',  'Identificador único'),
    ('id_empresa',            'INTEGER FK → empresas',     'No',  'Empresa asociada'),
    ('id_contacto',           'INTEGER FK → contactos',    'No',  'Contacto asociado'),
    ('puesto',                'TEXT',                       'Sí',  'Puesto del contacto'),
    ('categoria',             'TEXT',                       'Sí',  'Categoría del contacto'),
    ('id_direccion',          'INTEGER',                   'Sí',  'Referencia a dirección'),
    ('id_telefono',           'INTEGER',                   'Sí',  'Referencia a teléfono'),
    ('id_email',              'INTEGER',                   'Sí',  'Referencia a email'),
    ('es_usuario',            'INTEGER',                   'Sí',  '1 si es usuario del sistema'),
    ('created_at',            'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',            'DATETIME',                  'Sí',  'Última modificación'),
])

# es (estaciones de servicio)
db_table('es (Estaciones de Servicio)', [
    ('id_es',                 'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('id_empresa',            'INTEGER FK → empresas',    'Sí',  'Empresa propietaria'),
    ('cod_es',                'TEXT',                     'Sí',  'Código de la estación'),
    ('cod_retailgas',         'TEXT',                     'Sí',  'Código Retailgas'),
    ('cod_sociedad',          'TEXT',                     'Sí',  'Código de sociedad'),
    ('cod_solred',            'TEXT',                     'Sí',  'Código SOLRED'),
    ('concesion',             'TEXT',                     'Sí',  'Concesión'),
    ('tipo',                  'TEXT',                     'Sí',  'Tipo de estación'),
    ('nombre',                'TEXT',                     'Sí',  'Nombre de la estación'),
    ('direccion',             'TEXT',                     'Sí',  'Dirección completa'),
    ('cod_postal',            'TEXT',                     'Sí',  'Código postal'),
    ('poblacion',             'TEXT',                     'Sí',  'Población'),
    ('provincia',             'TEXT',                     'Sí',  'Provincia'),
    ('ccaa',                  'TEXT',                     'Sí',  'Comunidad autónoma'),
    ('pais',                  'TEXT',                     'Sí',  'País (def: España)'),
    ('num_margenes',          'INTEGER',                  'Sí',  'Número de márgenes'),
    ('y_wgs84',               'REAL',                     'Sí',  'Coordenada Y (latitud)'),
    ('x_wgs84',               'REAL',                     'Sí',  'Coordenada X (longitud)'),
    ('vinculo',               'TEXT',                     'Sí',  'Vínculo'),
    ('vinculo_2',             'TEXT',                     'Sí',  'Vínculo secundario'),
    ('delegacion',            'TEXT',                     'Sí',  'Delegación'),
    ('delegado',              'TEXT',                     'Sí',  'Delegado asignado'),
    ('tecnico_gestion',       'TEXT',                     'Sí',  'Técnico de gestión'),
    ('tl_tecnico_gestion',    'TEXT',                     'Sí',  'Teléfono técnico gestión'),
    ('email_tecnico_gestion', 'TEXT',                     'Sí',  'Email técnico gestión'),
    ('responsable_gestor',    'TEXT',                     'Sí',  'Responsable gestor'),
    ('tel_movil',             'TEXT',                     'Sí',  'Teléfono móvil'),
    ('tl_oficina',            'TEXT',                     'Sí',  'Teléfono oficina'),
    ('sede_email',            'TEXT',                     'Sí',  'Email de la sede'),
    ('tipo_mantenimiento',    'TEXT',                     'Sí',  'Tipo de mantenimiento'),
    ('f_alta',                'DATE',                     'Sí',  'Fecha de alta'),
    ('f_baja',                'DATE',                     'Sí',  'Fecha de baja'),
    ('f_alta_modificacion',   'DATE',                     'Sí',  'Fecha modificación alta'),
    ('nif',                   'TEXT',                     'Sí',  'NIF de la estación'),
    ('horario_apertura',      'TEXT',                     'Sí',  'Horario de apertura'),
    ('created_at',            'DATETIME',                 'Sí',  'Fecha de creación'),
    ('updated_at',            'DATETIME',                 'Sí',  'Última modificación'),
])

# contratos
db_table('contratos', [
    ('id_contrato',    'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('id_empresa',     'INTEGER FK → empresas',    'Sí',  'Empresa propietaria del contrato'),
    ('numero_contrato','VARCHAR',                   'Sí',  'Número de contrato'),
    ('nombre',         'VARCHAR',                   'Sí',  'Nombre descriptivo del contrato'),
    ('id_tarifario',   'INTEGER',                   'Sí',  'Tarifario vinculado al contrato'),
    ('fecha_inicio',   'DATE',                      'Sí',  'Fecha de inicio de vigencia'),
    ('fecha_fin',      'DATE',                      'Sí',  'Fecha de fin de vigencia'),
    ('descripcion',    'TEXT',                      'Sí',  'Descripción libre'),
    ('created_at',     'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',     'DATETIME',                  'Sí',  'Última modificación'),
])

# tarifario
db_table('tarifario', [
    ('id_tarifario',    'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('nombre_tarifario','VARCHAR',                   'No',  'Nombre del tarifario'),
    ('id_empresa',      'INTEGER',                   'Sí',  'Empresa asociada'),
    ('fecha_tarifario', 'DATE',                      'Sí',  'Fecha de inicio de vigencia'),
    ('fecha_fin',       'DATE',                      'Sí',  'Fecha de fin de vigencia'),
    ('created_at',      'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',      'DATETIME',                  'Sí',  'Última modificación'),
])

# tarifario_servicios
db_table('tarifario_servicios', [
    ('id_tarifario',   'INTEGER PK FK → tarifario', 'No',  'Tarifario al que pertenece'),
    ('id_servicio',    'INTEGER PK',                 'No',  'ID de servicio dentro del tarifario'),
    ('codigo_servicio','VARCHAR',                     'Sí',  'Código del servicio'),
    ('numero_tarifa',  'VARCHAR',                     'Sí',  'Número de tarifa'),
    ('nombre_servicio','VARCHAR',                     'No',  'Nombre del servicio'),
    ('precio_unitario','NUMERIC',                     'No',  'Precio por unidad'),
    ('precio_anterior','NUMERIC',                     'Sí',  'Precio anterior (histórico)'),
    ('created_at',     'DATETIME',                    'Sí',  'Fecha de creación'),
    ('updated_at',     'DATETIME',                    'Sí',  'Última modificación'),
])

# facturas
db_table('facturas', [
    ('id_factura',     'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('numero_factura', 'VARCHAR',                   'Sí',  'Número de factura'),
    ('id_proyecto',    'INTEGER',                   'Sí',  'Referencia a proyecto'),
    ('id_empresa',     'INTEGER',                   'Sí',  'Empresa facturada'),
    ('id_es',          'INTEGER',                   'Sí',  'Estación de servicio'),
    ('id_contrato',    'INTEGER',                   'Sí',  'Contrato asociado'),
    ('id_tarifario',   'INTEGER',                   'Sí',  'Tarifario aplicado'),
    ('id_direccion',   'INTEGER',                   'Sí',  'Dirección de facturación'),
    ('fecha_factura',  'DATE',                      'Sí',  'Fecha de la factura'),
    ('fecha_solicitud','DATE',                      'Sí',  'Fecha de solicitud'),
    ('importe_total',  'NUMERIC',                   'Sí',  'Importe total calculado'),
    ('created_at',     'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',     'DATETIME',                  'Sí',  'Última modificación'),
])

# factura_pedidos
db_table('factura_pedidos', [
    ('id_factura_pedido', 'INTEGER PK AUTOINCREMENT',  'No',  'Identificador único'),
    ('id_factura',        'INTEGER FK → facturas',     'No',  'Factura a la que pertenece'),
    ('id_pedido',         'INTEGER',                   'No',  'Pedido vinculado'),
    ('created_at',        'DATETIME',                  'Sí',  'Fecha de creación'),
    ('updated_at',        'DATETIME',                  'Sí',  'Última modificación'),
])

# lineas_factura
db_table('lineas_factura', [
    ('id_linea_factura', 'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('id_factura',       'INTEGER FK → facturas',    'No',  'Factura a la que pertenece'),
    ('id_servicio',      'INTEGER',                  'Sí',  'Servicio del tarifario'),
    ('id_tarifario',     'INTEGER',                  'Sí',  'Tarifario del servicio'),
    ('unidades',         'NUMERIC',                  'Sí',  'Cantidad de unidades'),
    ('precio_unitario',  'NUMERIC',                  'Sí',  'Precio unitario (del tarifario)'),
    ('importe',          'NUMERIC',                  'Sí',  'Importe = unidades × precio'),
    ('created_at',       'DATETIME',                 'Sí',  'Fecha de creación'),
    ('updated_at',       'DATETIME',                 'Sí',  'Última modificación'),
])

# cobros
db_table('cobros', [
    ('id_cobro',        'INTEGER PK AUTOINCREMENT', 'No',  'Identificador único'),
    ('id_factura',      'INTEGER FK → facturas',    'No',  'Factura cobrada'),
    ('importe',         'NUMERIC',                  'No',  'Importe del cobro'),
    ('fecha',           'DATE',                     'No',  'Fecha del cobro'),
    ('tipologia_cobro', 'VARCHAR',                  'Sí',  'Tipo de cobro (transferencia, etc.)'),
    ('cuenta_bancaria', 'VARCHAR',                  'Sí',  'Cuenta bancaria receptora'),
    ('created_at',      'DATETIME',                 'Sí',  'Fecha de creación'),
    ('updated_at',      'DATETIME',                 'Sí',  'Última modificación'),
])

h2('Relaciones entre Tablas')
code_block(
    'users\n'
    '  (contiene campo permisos JSON con acceso por módulo)\n'
    '\n'
    'presupuestos\n'
    '  └─ presupuestos_lineas  (id_presupuesto → presupuestos.id_presupuesto)\n'
    '  └─ pedidos              (id_presupuesto → presupuestos.id_presupuesto)\n'
    '\n'
    'pedidos\n'
    '  └─ lineas_pedido        (id_pedido → pedidos.id_pedido)\n'
    '  └─ factura_pedidos      (id_pedido → pedidos.id_pedido)\n'
    '\n'
    'empresas\n'
    '  └─ contactos_empresas   (id_empresa → empresas.id_empresa)\n'
    '  └─ es                   (id_empresa → empresas.id_empresa)\n'
    '  └─ contratos            (id_empresa → empresas.id_empresa)\n'
    '\n'
    'contactos\n'
    '  └─ contactos_empresas   (id_contacto → contactos.id_contacto)\n'
    '\n'
    'tarifario\n'
    '  └─ tarifario_servicios  (id_tarifario → tarifario.id_tarifario)\n'
    '  └─ contratos            (id_tarifario referencia no-FK)\n'
    '\n'
    'facturas\n'
    '  └─ factura_pedidos      (id_factura → facturas.id_factura)\n'
    '  └─ lineas_factura       (id_factura → facturas.id_factura)\n'
    '  └─ cobros               (id_factura → facturas.id_factura)'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. GUÍA DE INSTALACIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Guía de Instalación')

h2('3.1  Requisitos previos')
body('Instala las siguientes herramientas antes de continuar:')

h3('Python 3.11 o superior')
bullet('Descarga: https://www.python.org/downloads/')
bullet('Marca la opción "Add Python to PATH" durante la instalación.')
bullet('Verifica en terminal: python --version')

h3('Node.js 18 o superior (incluye npm)')
bullet('Descarga: https://nodejs.org/  (versión LTS recomendada)')
bullet('Verifica en terminal: node --version  y  npm --version')

h3('Angular CLI')
code_block('npm install -g @angular/cli')
bullet('Verifica: ng version')

h3('(Opcional) Visual Studio Code')
bullet('Descarga: https://code.visualstudio.com/')
bullet('Extensiones recomendadas: Python (Microsoft), Angular Language Service, '
       'SQLite Viewer (Florian Klampfer).')

doc.add_paragraph()
h2('3.2  Instalación del Backend')
body('Abre una terminal (PowerShell o CMD) y ejecuta:')
code_block(
    '# 1. Entra en la carpeta del backend\n'
    'cd C:\\Users\\Bernardo\\Documents\\erp-ciete\\erp-ciete\\erp-ciete\\backend\n'
    '\n'
    '# 2. Instala las dependencias\n'
    'python -m pip install -r requirements.txt\n'
    '\n'
    '# 3. Comprueba que no hay errores'
)
body('Si la instalación es correcta verás "Successfully installed ..." sin errores en rojo.')
body(
    'NOTA: No es necesario crear la base de datos manualmente. '
    'El servidor la crea automáticamente en el primer arranque. '
    'Si la BD ya existía de una versión anterior, el sistema añade automáticamente '
    'las nuevas columnas y tablas sin perder los datos existentes.'
)

doc.add_paragraph()
h2('3.3  Instalación del Frontend')
body('Abre otra terminal y ejecuta:')
code_block(
    '# 1. Entra en la carpeta del frontend\n'
    'cd C:\\Users\\Bernardo\\Documents\\erp-ciete\\erp-ciete\\erp-ciete\\frontend\n'
    '\n'
    '# 2. Instala los paquetes de Node\n'
    'npm install\n'
    '\n'
    '# (puede tardar 1-2 minutos la primera vez)'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. CREDENCIALES
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Credenciales de Acceso')

h2('4.1  Usuario Administrador por Defecto')
body(
    'Al arrancar el backend por primera vez se crea automáticamente '
    'un usuario administrador con las siguientes credenciales:'
)
t = doc.add_table(rows=1, cols=2)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Campo', 'Valor']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '1E3A8A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for k, v in [
    ('Email / Usuario', 'admin@ciete.es'),
    ('Contraseña',      'admin123'),
    ('Rol',             'admin'),
    ('Nombre',          'Admin'),
    ('Permisos',        'Todos los módulos activados'),
]:
    row = t.add_row().cells
    row[0].text = k
    row[1].text = v

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    '⚠  IMPORTANTE: Cambia la contraseña del administrador en producción. '
    'Puedes hacerlo desde la sección Usuarios → Editar o directamente en la '
    'base de datos con DB Browser for SQLite.'
)
run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
run.font.bold = True

doc.add_paragraph()
h2('4.2  Clave Secreta JWT')
body(
    'La clave secreta usada para firmar los tokens JWT está definida en '
    'backend/auth.py en la constante SECRET_KEY. '
    'En un entorno de producción debe moverse a una variable de entorno.'
)
code_block(
    'SECRET_KEY = "4f3e2a1b9c8d7e6f5a4b3c2d1e0f9a8b7c6d5e4f3a2b1c0d9e8f7a6b5c4d3e2"\n'
    'ALGORITHM  = "HS256"\n'
    'Expiración = 8 horas'
)

doc.add_paragraph()
h2('4.3  URLs de Acceso')
t = doc.add_table(rows=1, cols=2)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Servicio', 'URL']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '1E3A8A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for s, u in [
    ('Aplicación web (frontend)',     'http://localhost:4200'),
    ('API REST (backend)',            'http://localhost:8000'),
    ('Documentación Swagger (API)',   'http://localhost:8000/docs'),
    ('Documentación ReDoc (API)',     'http://localhost:8000/redoc'),
]:
    row = t.add_row().cells
    row[0].text = s
    row[1].text = u

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. GUÍA DE USO DIARIO
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Guía de Uso Diario')

h2('5.1  Arrancar el Sistema (cada vez que quieras usar el ERP)')
body(
    'Para que la aplicación funcione necesitas tener DOS terminales abiertas '
    'de forma simultánea: una para el backend y otra para el frontend.'
)

h3('Terminal 1 — Backend')
code_block(
    'cd C:\\Users\\Bernardo\\Documents\\erp-ciete\\erp-ciete\\erp-ciete\\backend\n'
    'uvicorn main:app --reload'
)
body('Debes ver algo similar a:')
code_block(
    'INFO:     Uvicorn running on http://127.0.0.1:8000 (Press CTRL+C to quit)\n'
    'INFO:     Started reloader process [...]\n'
    'INFO:     Application startup complete.'
)
bullet('El flag --reload hace que el servidor se reinicie automáticamente al editar archivos Python.')
bullet('Deja esta terminal abierta mientras usas el ERP.')

doc.add_paragraph()
h3('Terminal 2 — Frontend')
code_block(
    'cd C:\\Users\\Bernardo\\Documents\\erp-ciete\\erp-ciete\\erp-ciete\\frontend\n'
    'ng serve'
)
body('Debes ver algo similar a:')
code_block(
    '✔ Compiled successfully.\n'
    'Watch mode enabled. Watching for file changes...\n'
    '  ➜  Local:   http://localhost:4200/'
)
bullet('Abre http://localhost:4200 en el navegador.')
bullet('El frontend se recompila automáticamente cuando editas archivos TypeScript/HTML/CSS.')
bullet('Deja esta terminal abierta mientras usas el ERP.')

doc.add_paragraph()
h2('5.2  Acceder a la Aplicación')
body('Abre el navegador y ve a: http://localhost:4200')
bullet('Introduce el email (o identificador) y contraseña del administrador.')
bullet('Al hacer login, tu sesión se guarda en el navegador (localStorage) y '
       'dura 8 horas antes de expirar.')
bullet('El dashboard mostrará solo las tarjetas de módulos a los que tienes acceso.')

doc.add_paragraph()
h2('5.3  Flujo Principal: Presupuesto → Pedido → Factura')
steps = [
    ('1', 'Ir a Presupuestos',   'Haz clic en "Presupuestos" en el menú lateral.'),
    ('2', 'Nuevo Presupuesto',   'Pulsa "+ Nuevo Presupuesto", elige empresa, tarifario, '
                                  'fecha y añade líneas (servicio + unidades). Pulsa "Guardar".'),
    ('3', 'Enviar',              'En el detalle del presupuesto, cambia el estado a "Enviado" '
                                  'usando los botones de estado.'),
    ('4', 'Aprobar',             'Cuando el cliente acepte, cambia el estado a "Aprobado".'),
    ('5', 'Convertir a Pedido',  'Con estado "Aprobado" aparece el botón verde '
                                  '"Convertir a Pedido". Al pulsarlo se crea automáticamente '
                                  'el pedido copiando todos los datos y líneas.'),
    ('6', 'Gestionar el Pedido', 'Ve a la sección Pedidos para ver y gestionar el estado '
                                  'del pedido generado.'),
    ('7', 'Crear Factura',       'Ve a Facturas → Nueva Factura. Selecciona empresa, '
                                  'proyecto, contrato, vincula el pedido, añade líneas '
                                  'del tarifario y guarda.'),
    ('8', 'Registrar Cobro',     'En el detalle de la factura, registra el cobro indicando '
                                  'importe, fecha y tipo. El estado de cobro se calcula '
                                  'automáticamente (Pendiente / Cobro parcial / Cobrada).'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Paso', 'Acción', 'Detalle']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '16A34A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for num, acc, det in steps:
    row = t.add_row().cells
    row[0].text = num
    row[1].text = acc
    row[2].text = det

doc.add_paragraph()
h2('5.4  Parar el Sistema')
bullet('En cada terminal pulsa Ctrl + C para detener el servidor.')
bullet('La base de datos se guarda automáticamente; no hay que hacer nada adicional.')

doc.add_paragraph()
h2('5.5  Preguntas Frecuentes')
faqs = [
    ('¿Se pierden los datos al parar?',
     'No. Los datos se guardan en backend/database.sqlite, que persiste entre reinicios.'),
    ('¿Cómo añadir nuevos usuarios?',
     'Ve a Usuarios (visible solo para administradores) y pulsa "+ Nuevo Usuario". '
     'Rellena nombre, identificador, contraseña, rol y permisos.'),
    ('¿Qué diferencia hay entre role=admin y role=user?',
     'El administrador puede gestionar usuarios y tiene acceso a todo. '
     'Un usuario normal solo accede a los módulos que el admin haya habilitado '
     'en su perfil de permisos.'),
    ('El login dice "Credenciales incorrectas"',
     'Asegúrate de que el backend está corriendo en http://localhost:8000 '
     'y de que la terminal no muestra errores.'),
    ('¿Cómo ver la base de datos?',
     'Descarga "DB Browser for SQLite" (https://sqlitebrowser.org/) '
     'y abre el archivo backend/database.sqlite para explorar y editar tablas.'),
    ('¿Dónde cambiar la URL del backend?',
     'En frontend/src/environments/environment.ts, variable apiUrl.'),
    ('No veo un módulo en el menú',
     'El módulo puede estar deshabilitado en tus permisos de usuario. '
     'Pide al administrador que los active desde Usuarios → Editar.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'P: {q}')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p2 = doc.add_paragraph(f'R: {a}')
    p2.paragraph_format.left_indent = Inches(0.3)
    doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MÓDULO DE EMPRESAS
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Módulo de Empresas')

body(
    'El módulo de Empresas permite gestionar las empresas registradas en el sistema. '
    'Cada empresa se identifica por su CIF y puede tener contactos asociados a través '
    'de la tabla intermedia contactos_empresas.'
)

h2('6.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Empresas".')
bullet('También puedes acceder desde el menú de navegación superior → Empresas.')

h2('6.2  Listado de Empresas')
bullet('Se muestra una tabla con ID, Nombre, Razón Social y CIF.')
bullet('Desde la lista puedes Ver el detalle, Editar o Eliminar cada empresa.')

h2('6.3  Crear una Empresa')
bullet('Pulsa el botón "+ Nueva Empresa".')
bullet('Rellena el formulario: Nombre (obligatorio), Razón Social y CIF.')
bullet('Pulsa "Guardar Empresa" para registrarla.')

h2('6.4  Editar una Empresa')
bullet('Desde el listado o el detalle, pulsa "Editar".')
bullet('Modifica los campos necesarios y pulsa "Actualizar".')

h2('6.5  Eliminar una Empresa')
bullet('Desde el listado, pulsa "Eliminar" en la fila correspondiente.')
bullet('Confirma la eliminación en el diálogo emergente.')
bullet('⚠ Se eliminarán también las relaciones en contactos_empresas y estaciones asociadas.')

h2('6.6  Endpoints API')
make_api_table([
    ('GET',    '/empresas',               'Lista todas las empresas'),
    ('POST',   '/empresas',               'Crea una nueva empresa'),
    ('GET',    '/empresas/{id}',           'Detalle de una empresa'),
    ('PUT',    '/empresas/{id}',           'Actualiza una empresa'),
    ('DELETE', '/empresas/{id}',           'Elimina una empresa'),
    ('GET',    '/empresas/{id}/contactos', 'Contactos de una empresa'),
    ('POST',   '/empresas/{id}/contactos', 'Vincula contacto a empresa'),
], 'D97706')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. MÓDULO DE CONTACTOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('7. Módulo de Contactos')

body(
    'El módulo de Contactos permite gestionar las personas de contacto. '
    'Un contacto puede vincularse a una o varias empresas a través de la '
    'tabla contactos_empresas, donde se registra su puesto y categoría.'
)

h2('7.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Contactos".')
bullet('También puedes acceder desde el menú de navegación superior → Contactos.')

h2('7.2  Listado de Contactos')
bullet('Se muestra una tabla con ID, Nombre, Primer Apellido y Segundo Apellido.')
bullet('Desde la lista puedes Ver, Editar o Eliminar cada contacto.')

h2('7.3  Crear un Contacto')
bullet('Pulsa el botón "+ Nuevo Contacto".')
bullet('Rellena el formulario: Nombre (obligatorio), Primer Apellido y Segundo Apellido.')
bullet('Pulsa "Guardar Contacto" para registrarlo.')

h2('7.4  Editar un Contacto')
bullet('Desde el listado o el detalle, pulsa "Editar".')
bullet('Modifica los campos necesarios y pulsa "Actualizar".')

h2('7.5  Eliminar un Contacto')
bullet('Desde el listado, pulsa "Eliminar" en la fila correspondiente.')
bullet('Confirma la eliminación en el diálogo emergente.')

h2('7.6  Endpoints API')
make_api_table([
    ('GET',    '/contactos',       'Lista todos los contactos'),
    ('POST',   '/contactos',       'Crea un nuevo contacto'),
    ('GET',    '/contactos/{id}',  'Detalle de un contacto'),
    ('PUT',    '/contactos/{id}',  'Actualiza un contacto'),
    ('DELETE', '/contactos/{id}',  'Elimina un contacto'),
], '7C3AED')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. MÓDULO DE ESTACIONES DE SERVICIO
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. Módulo de Estaciones de Servicio')

body(
    'El módulo de Estaciones de Servicio (ES) gestiona la información completa de cada '
    'estación: identificación, códigos, ubicación geográfica, datos de gestión, '
    'contacto y fechas operativas. Cada estación puede estar vinculada a una empresa.'
)

h2('8.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Estaciones de Servicio".')
bullet('También puedes acceder desde el menú de navegación superior → Estaciones.')

h2('8.2  Listado de Estaciones')
bullet('Se muestra una tabla con ID, Código ES, Nombre, Población, Provincia y Tipo.')
bullet('Desde la lista puedes Ver, Editar o Eliminar cada estación.')

h2('8.3  Crear una Estación')
bullet('Pulsa el botón "+ Nueva Estación".')
bullet('El formulario se organiza en secciones: Identificación, Códigos, Ubicación, '
       'Gestión, Contacto y Fechas.')
bullet('Rellena los campos necesarios y pulsa "Guardar Estación".')

h2('8.4  Editar una Estación')
bullet('Desde el listado o el detalle, pulsa "Editar".')
bullet('Modifica los campos necesarios y pulsa "Actualizar".')

h2('8.5  Eliminar una Estación')
bullet('Desde el listado, pulsa "Eliminar" en la fila correspondiente.')
bullet('Confirma la eliminación en el diálogo emergente.')

h2('8.6  Campos de la Estación')
body('La ficha de una estación contiene los siguientes grupos de información:')
bullet('Identificación: nombre, código ES, tipo, NIF, concesión, empresa asociada.')
bullet('Códigos: Retailgas, sociedad, SOLRED, vínculos.')
bullet('Ubicación: dirección, CP, población, provincia, CCAA, país, coordenadas.')
bullet('Gestión: delegación, delegado, técnico de gestión, responsable, mantenimiento.')
bullet('Contacto: teléfono móvil, teléfono oficina, email de la sede.')
bullet('Fechas: alta, baja, modificación de alta, horario de apertura.')

h2('8.7  Endpoints API')
make_api_table([
    ('GET',    '/estaciones',       'Lista todas las estaciones'),
    ('POST',   '/estaciones',       'Crea una nueva estación'),
    ('GET',    '/estaciones/{id}',  'Detalle de una estación'),
    ('PUT',    '/estaciones/{id}',  'Actualiza una estación'),
    ('DELETE', '/estaciones/{id}',  'Elimina una estación'),
], 'DC2626')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. MÓDULO DE PRESUPUESTOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('9. Módulo de Presupuestos')

body(
    'El módulo de Presupuestos permite crear ofertas económicas detalladas con líneas '
    'de servicio vinculadas a un tarifario. Incluye un flujo de estado y la posibilidad '
    'de convertir presupuestos aprobados en pedidos automáticamente.'
)

h2('9.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Presupuestos".')
bullet('También puedes acceder desde el menú de navegación superior → Presupuestos.')

h2('9.2  Listado de Presupuestos')
bullet('Se muestra una tabla con ID, Fecha, Estado y datos de proyecto/empresa.')
bullet('Desde la lista puedes Ver el detalle o acceder al formulario.')

h2('9.3  Crear un Presupuesto')
bullet('Pulsa el botón "+ Nuevo Presupuesto".')
bullet('Selecciona empresa, estación de servicio, tarifario y fecha.')
bullet('Añade líneas: elige servicio del tarifario y especifica las unidades.')
bullet('El sistema muestra el precio unitario del tarifario seleccionado.')
bullet('Pulsa "Guardar" para registrar el presupuesto con estado "borrador".')

h2('9.4  Flujo de Estados')
body('Los presupuestos siguen este flujo de estado:')
t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Estado', 'Significado', 'Transición posible']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '16A34A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for est, sig, trans in [
    ('borrador',  'Creado, pendiente de envío',        'Enviado'),
    ('enviado',   'Enviado al cliente',                'Aprobado / Rechazado'),
    ('aprobado',  'Aceptado por el cliente',           'Convertir a Pedido'),
    ('rechazado', 'Rechazado por el cliente',          'Final (no hay transición)'),
]:
    row = t.add_row().cells
    row[0].text = est
    row[1].text = sig
    row[2].text = trans

doc.add_paragraph()
h2('9.5  Detalle del Presupuesto')
bullet('Muestra todos los datos del presupuesto con sus líneas.')
bullet('Cada línea muestra: código de servicio, nombre, precio unitario, unidades e importe.')
bullet('Se calcula el total sumando todos los importes de las líneas.')
bullet('Botones de cambio de estado según el estado actual.')
bullet('Botón "Convertir a Pedido" visible cuando el estado es "aprobado".')

h2('9.6  Endpoints API')
make_api_table([
    ('GET',   '/presupuestos',                'Lista todos los presupuestos'),
    ('POST',  '/presupuestos',                'Crea presupuesto con líneas'),
    ('GET',   '/presupuestos/{id}',           'Detalle con líneas cargadas'),
    ('PATCH', '/presupuestos/{id}/estado',    'Cambia el estado del presupuesto'),
    ('POST',  '/presupuestos/{id}/convertir', 'Convierte a pedido (estado=aprobado)'),
], '16A34A')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. MÓDULO DE PEDIDOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('10. Módulo de Pedidos')

body(
    'El módulo de Pedidos gestiona los pedidos generados a partir de presupuestos aprobados. '
    'Cada pedido hereda los datos y líneas del presupuesto de origen y puede vincularse '
    'posteriormente a facturas.'
)

h2('10.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Pedidos".')
bullet('También puedes acceder desde el menú de navegación superior → Pedidos.')

h2('10.2  Listado de Pedidos')
bullet('Se muestra una tabla con ID, Estado, Empresa, Fecha de solicitud y Proyecto.')
bullet('Desde la lista puedes acceder al detalle completo de cada pedido.')

h2('10.3  Detalle del Pedido')
bullet('Muestra todos los datos del pedido: empresa, proyecto, estación, tarifario.')
bullet('Lista las líneas de pedido con servicio y unidades.')
bullet('Muestra el presupuesto de origen vinculado.')

h2('10.4  Endpoints API')
make_api_table([
    ('GET', '/pedidos',       'Lista todos los pedidos'),
    ('GET', '/pedidos/{id}',  'Detalle completo de un pedido'),
], '0891B2')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. MÓDULO DE CONTRATOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('11. Módulo de Contratos')

body(
    'El módulo de Contratos gestiona los acuerdos comerciales entre Ciete Ingenieros '
    'y sus clientes. Cada contrato está vinculado a una empresa y puede tener un '
    'tarifario asociado que define los precios aplicables. El sistema calcula '
    'automáticamente si un contrato está vigente o caducado según su fecha de fin.'
)

h2('11.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Contratos" (si tienes permiso).')
bullet('El menú de navegación superior → Contratos (visible si el permiso está activo).')

h2('11.2  Listado de Contratos')
bullet('Tabla con: Número de contrato, Empresa, Tarifario vinculado, Fecha inicio, Fecha fin.')
bullet('Estado calculado automáticamente: Vigente (fecha_fin futura) o Caducado (fecha_fin pasada).')
bullet('Botones de Editar y Eliminar por fila.')

h2('11.3  Crear un Contrato')
bullet('Pulsa el botón "+ Nuevo Contrato".')
body('El formulario contiene los siguientes campos:')
bullet('Empresa: selector desplegable con todas las empresas registradas.')
bullet('Número de contrato: identificador único del contrato.')
bullet('Nombre: descripción breve del contrato.')
bullet('Tarifario: selector con los tarifarios disponibles.')
bullet('Fecha inicio: inicio de la vigencia del contrato.')
bullet('Fecha fin: fin de la vigencia del contrato.')
bullet('Descripción: texto libre con detalles adicionales.')
bullet('Pulsa "Guardar" para registrar el contrato.')

h2('11.4  Editar un Contrato')
bullet('Desde el listado, pulsa el botón "Editar" en la fila del contrato.')
bullet('Modifica los campos necesarios y pulsa "Actualizar".')

h2('11.5  Eliminar un Contrato')
bullet('Desde el listado, pulsa "Eliminar".')
bullet('Confirma la acción en el diálogo emergente.')

h2('11.6  Permisos requeridos')
body(
    'El módulo de Contratos requiere que el usuario tenga el permiso "contratos" activado '
    'en su perfil. Los administradores tienen acceso completo por defecto.'
)

h2('11.7  Endpoints API')
make_api_table([
    ('GET',    '/contratos',       'Lista todos los contratos'),
    ('POST',   '/contratos',       'Crea un nuevo contrato'),
    ('GET',    '/contratos/{id}',  'Detalle de contrato con empresa y tarifario'),
    ('PUT',    '/contratos/{id}',  'Actualiza un contrato'),
    ('DELETE', '/contratos/{id}',  'Elimina un contrato'),
], '059669')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. MÓDULO DE TARIFARIO
# ═══════════════════════════════════════════════════════════════════════════════
h1('12. Módulo de Tarifario')

body(
    'El módulo de Tarifario gestiona las listas de precios de servicios. '
    'Cada tarifario puede tener múltiples servicios con su código, nombre y precio unitario. '
    'Los tarifarios se vinculan a contratos, presupuestos y facturas para aplicar '
    'los precios correctos.'
)

h2('12.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Tarifario" (si tienes permiso).')
bullet('El menú de navegación superior → Tarifario.')

h2('12.2  Listado de Tarifarios')
bullet('Tabla con: Nombre, Empresa asociada, Fecha de inicio (fecha_tarifario), Fecha de fin.')
bullet('Pulsar sobre un tarifario abre su vista de detalle.')

h2('12.3  Crear un Tarifario')
bullet('Pulsa el botón "+ Nuevo Tarifario".')
body('Campos del formulario:')
bullet('Nombre del tarifario.')
bullet('Empresa: selector opcional.')
bullet('Fecha de inicio de vigencia.')
bullet('Fecha de fin de vigencia.')
bullet('Pulsa "Guardar" para crear el tarifario (inicialmente sin servicios).')

h2('12.4  Gestión de Servicios del Tarifario')
body(
    'La vista de detalle del tarifario muestra una tabla editable con todos sus servicios. '
    'Desde esta misma pantalla puedes:'
)
bullet('Añadir un nuevo servicio: rellena el formulario inline con código, número de tarifa, '
       'nombre y precio unitario, y pulsa "Añadir".')
bullet('Editar un servicio existente: pulsa el icono de edición en la fila para activar '
       'los campos editables; confirma con "Guardar".')
bullet('Eliminar un servicio: pulsa el icono de papelera y confirma.')

h2('12.5  Campos de un Servicio')
t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Campo', 'Descripción', 'Obligatorio']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '2563EB')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for campo, desc, req in [
    ('codigo_servicio', 'Código alfanumérico del servicio',          'No'),
    ('numero_tarifa',   'Número de tarifa interno',                  'No'),
    ('nombre_servicio', 'Nombre descriptivo del servicio',           'Sí'),
    ('precio_unitario', 'Precio por unidad (€)',                     'Sí'),
    ('precio_anterior', 'Precio histórico anterior (referencia)',     'No'),
]:
    row = t.add_row().cells
    row[0].text = campo
    row[1].text = desc
    row[2].text = req

doc.add_paragraph()
h2('12.6  Permisos requeridos')
body(
    'El módulo de Tarifario requiere que el usuario tenga el permiso "tarifario" activado. '
    'Los administradores tienen acceso completo por defecto.'
)

h2('12.7  Endpoints API')
make_api_table([
    ('GET',    '/tarifario',                         'Lista todos los tarifarios'),
    ('POST',   '/tarifario',                         'Crea un nuevo tarifario'),
    ('GET',    '/tarifario/{id}',                    'Detalle con sus servicios'),
    ('PUT',    '/tarifario/{id}',                    'Actualiza un tarifario'),
    ('DELETE', '/tarifario/{id}',                    'Elimina tarifario y sus servicios'),
    ('GET',    '/tarifario/{id}/servicios',          'Lista servicios del tarifario'),
    ('POST',   '/tarifario/{id}/servicios',          'Añade un servicio'),
    ('PUT',    '/tarifario/{id_t}/servicios/{id_s}', 'Actualiza un servicio'),
    ('DELETE', '/tarifario/{id_t}/servicios/{id_s}', 'Elimina un servicio'),
], '2563EB')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. MÓDULO DE FACTURAS
# ═══════════════════════════════════════════════════════════════════════════════
h1('13. Módulo de Facturas')

body(
    'El módulo de Facturas permite crear y gestionar las facturas emitidas. '
    'Cada factura puede vincularse a pedidos existentes, contener líneas de '
    'servicios del tarifario y registrar cobros. El sistema calcula automáticamente '
    'el estado de cobro en función de los cobros registrados.'
)

h2('13.1  Acceso')
bullet('Desde el Dashboard, haz clic en la tarjeta "Facturas" (si tienes permiso).')
bullet('El menú de navegación superior → Facturas.')

h2('13.2  Listado de Facturas')
bullet('Tabla con: Número de factura, Empresa, Fecha, Importe total, Estado de cobro.')
body('El estado de cobro se calcula automáticamente:')
t = doc.add_table(rows=1, cols=2)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Estado', 'Condición']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '1E3A8A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for est, cond in [
    ('Pendiente',     'No hay cobros registrados'),
    ('Cobro parcial', 'La suma de cobros es menor que el importe total'),
    ('Cobrada',       'La suma de cobros es igual o mayor al importe total'),
]:
    row = t.add_row().cells
    row[0].text = est
    row[1].text = cond

doc.add_paragraph()
h2('13.3  Crear una Factura')
bullet('Pulsa el botón "+ Nueva Factura".')
body('El formulario se divide en secciones:')
h3('Datos generales')
bullet('Número de factura.')
bullet('Empresa: selector desplegable.')
bullet('Proyecto: referencia numérica de proyecto.')
bullet('Estación de servicio: selector opcional.')
bullet('Contrato: selector filtrado por la empresa seleccionada.')
bullet('Fecha de factura y fecha de solicitud.')

h3('Pedidos vinculados')
bullet('Permite buscar y vincular uno o varios pedidos existentes a la factura.')
bullet('Se puede añadir múltiples pedidos para agrupar trabajo en una sola factura.')

h3('Líneas de factura')
bullet('Tabla editable donde cada línea tiene:')
bullet('  — Selector de tarifario.')
bullet('  — Selector de servicio (se carga al elegir tarifario, con precio unitario).')
bullet('  — Unidades: al escribirlas el importe se calcula automáticamente (precio × unidades).')
bullet('  — Importe calculado (solo lectura).')

h2('13.4  Detalle de una Factura')
bullet('Muestra todos los datos generales, pedidos vinculados y líneas con sus importes.')
bullet('Suma total al final de las líneas.')
body('Sección de cobros:')
bullet('Lista todos los cobros registrados con fecha, importe y tipo.')
bullet('Formulario inline para registrar un nuevo cobro sin salir de la pantalla.')
bullet('Botón de eliminar cobro por fila.')

h2('13.5  Registrar un Cobro')
body('Desde el detalle de la factura, en la sección de cobros:')
bullet('Introduce el importe del cobro.')
bullet('Selecciona la fecha del cobro.')
bullet('Indica la tipología de cobro (transferencia, cheque, etc.).')
bullet('Indica la cuenta bancaria receptora.')
bullet('Pulsa "Registrar Cobro".')

h2('13.6  Permisos requeridos')
body(
    'El módulo de Facturas requiere que el usuario tenga el permiso "facturas" activado. '
    'Los administradores tienen acceso completo por defecto.'
)

h2('13.7  Endpoints API')
make_api_table([
    ('GET',    '/facturas',                        'Lista todas las facturas'),
    ('POST',   '/facturas',                        'Crea factura con líneas y pedidos'),
    ('GET',    '/facturas/{id}',                   'Detalle completo con cobros'),
    ('PUT',    '/facturas/{id}',                   'Actualiza datos generales'),
    ('DELETE', '/facturas/{id}',                   'Elimina factura y sus datos relacionados'),
    ('GET',    '/facturas/{id}/cobros',            'Lista cobros de la factura'),
    ('POST',   '/facturas/{id}/cobros',            'Registra un cobro'),
    ('DELETE', '/facturas/{id}/cobros/{id_cobro}', 'Elimina un cobro'),
], 'B45309')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. GESTIÓN DE USUARIOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('14. Gestión de Usuarios')

body(
    'El módulo de Usuarios permite al administrador del sistema crear, editar y eliminar '
    'cuentas de usuario. Solo los usuarios con rol "admin" pueden acceder a esta sección. '
    'El campo de login puede ser un email convencional o un identificador simple '
    '(el sistema no valida formato de email).'
)

h2('14.1  Acceso')
bullet('El enlace "Usuarios" en la barra de navegación solo es visible para usuarios con role=admin.')
bullet('URL directa: http://localhost:4200/usuarios')

h2('14.2  Listado de Usuarios')
bullet('Tabla con: Nombre, Email/Usuario, Rol y botones de Editar y Eliminar.')
bullet('La contraseña nunca se muestra en ninguna vista.')
bullet('No se puede eliminar el propio usuario que está logueado.')

h2('14.3  Crear un Usuario')
bullet('Pulsa el botón "+ Nuevo Usuario".')
body('El formulario contiene:')
bullet('Nombre: nombre completo del usuario.')
bullet('Email o identificador: puede ser un email (admin@ciete.es) o un identificador '
       'simple (operador1). No se valida formato.')
bullet('Contraseña: mínimo requerido. Se almacena hasheada con bcrypt.')
bullet('Rol: selector con opciones "user" o "admin".')
body('Sección de permisos por módulo (checkboxes o toggles):')
bullet('Presupuestos, Pedidos, Facturas, Contratos, Tarifario, Usuarios.')
bullet('Al seleccionar rol "admin" todos los permisos se activan automáticamente.')
bullet('Al seleccionar rol "user" los permisos quedan como estaban (se pueden ajustar manualmente).')

h2('14.4  Editar un Usuario')
bullet('Desde el listado, pulsa "Editar" en la fila del usuario.')
bullet('Puedes cambiar nombre, email/identificador, rol y permisos.')
bullet('El campo contraseña es opcional en edición: si se deja vacío, la contraseña actual no cambia.')
bullet('Si se introduce una nueva contraseña, se hashea con bcrypt antes de guardar.')
bullet('Pulsa "Guardar" para confirmar los cambios.')

h2('14.5  Eliminar un Usuario')
bullet('Desde el listado, pulsa "Eliminar" en la fila del usuario.')
bullet('Confirma la acción en el diálogo emergente.')
bullet('⚠ No puedes eliminar tu propio usuario (el sistema lo impide con un error 400).')

h2('14.6  Permisos por Módulo — Formato JSON')
body(
    'Los permisos se almacenan en la columna "permisos" de la tabla users como una '
    'cadena JSON con este formato:'
)
code_block(
    '{\n'
    '  "presupuestos": true,\n'
    '  "pedidos": true,\n'
    '  "facturas": true,\n'
    '  "contratos": true,\n'
    '  "tarifario": true,\n'
    '  "usuarios": false\n'
    '}'
)
body('Valores por defecto según rol:')
t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Módulo', 'admin (por defecto)', 'user (por defecto)']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '1E3A8A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for modulo, adm, usr in [
    ('presupuestos', 'true', 'false'),
    ('pedidos',      'true', 'false'),
    ('facturas',     'true', 'false'),
    ('contratos',    'true', 'false'),
    ('tarifario',    'true', 'false'),
    ('usuarios',     'true', 'false'),
]:
    row = t.add_row().cells
    row[0].text = modulo
    row[1].text = adm
    row[2].text = usr

doc.add_paragraph()
h2('14.7  Endpoints API (solo admin)')
make_api_table([
    ('GET',    '/usuarios',       'Lista todos los usuarios (sin contraseña)'),
    ('POST',   '/usuarios',       'Crea un nuevo usuario con rol y permisos'),
    ('GET',    '/usuarios/{id}',  'Detalle de un usuario'),
    ('PUT',    '/usuarios/{id}',  'Edita nombre, email, rol, permisos y contraseña opcional'),
    ('DELETE', '/usuarios/{id}',  'Elimina usuario (no el propio)'),
], '991B1B')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 15. SISTEMA DE PERMISOS Y ROLES
# ═══════════════════════════════════════════════════════════════════════════════
h1('15. Sistema de Permisos y Roles')

h2('15.1  Roles del Sistema')
body('El ERP Ciete Ingenieros tiene dos roles de usuario:')
t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Rol', 'Descripción', 'Acceso a Usuarios']):
    t.rows[0].cells[i].text = h
    shade_cell(t.rows[0].cells[i], '1E3A8A')
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for rol, desc, acc in [
    ('admin', 'Administrador del sistema. Acceso completo a todos los módulos '
               'y a la gestión de usuarios.', 'Sí (total)'),
    ('user',  'Usuario estándar. Solo accede a los módulos habilitados '
               'en su perfil de permisos.', 'No'),
]:
    row = t.add_row().cells
    row[0].text = rol
    row[1].text = desc
    row[2].text = acc

doc.add_paragraph()
h2('15.2  Guard de Autenticación (AuthGuard)')
body(
    'AuthGuard protege todas las rutas que requieren que el usuario esté logueado. '
    'Si el navegador no tiene un JWT válido en localStorage, redirige automáticamente '
    'a la página de login.'
)

h2('15.3  Guard de Permisos (PermisosGuard)')
body(
    'PermisosGuard protege los módulos nuevos (contratos, tarifario, facturas, usuarios). '
    'Al intentar acceder a una ruta protegida, el guard verifica:'
)
bullet('1. Que el usuario está autenticado (JWT válido).')
bullet('2. Que el campo permisos del usuario tiene el módulo correspondiente en true, '
       'o que el rol es "admin".')
bullet('Si no tiene permiso, redirige al dashboard con el mensaje: '
       '"No tienes acceso a este módulo".')

h2('15.4  Dashboard Filtrado por Permisos')
body(
    'El DashboardComponent muestra solo las tarjetas de los módulos a los que el '
    'usuario tiene acceso según su combinación de rol y permisos:'
)
bullet('Los módulos sin permiso no aparecen en el dashboard.')
bullet('El enlace "Usuarios" en la barra de navegación solo aparece si role=admin.')
bullet('Los usuarios admin ven todas las tarjetas del dashboard.')

h2('15.5  Flujo de Validación de Permisos')
code_block(
    'Usuario hace clic en "Contratos"\n'
    '  → AuthGuard: ¿hay JWT válido?  → No → Redirige a /login\n'
    '                                  → Sí ↓\n'
    '  → PermisosGuard: ¿role=admin? → Sí → Permite acceso\n'
    '                    ¿No admin?  → ¿permisos.contratos === true?\n'
    '                                     → Sí → Permite acceso\n'
    '                                     → No → Redirige a /dashboard\n'
    '                                             (mensaje: Sin acceso)'
)

doc.add_paragraph()
h2('15.6  Cómo Gestionar Permisos')
body('Solo un administrador puede modificar los permisos de un usuario:')
bullet('1. Ve a Usuarios (menú superior, solo visible para admin).')
bullet('2. Busca el usuario en la lista y pulsa "Editar".')
bullet('3. En la sección de permisos, activa o desactiva los toggles de cada módulo.')
bullet('4. Pulsa "Guardar" para aplicar los cambios.')
bullet('Los cambios son efectivos en el siguiente inicio de sesión del usuario afectado '
       '(o de inmediato si el token se refresca).')

# ═══════════════════════════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ═══════════════════════════════════════════════════════════════════════════════
section = doc.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'ERP Ciete Ingenieros · Documentación Técnica v3.0 · {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# GUARDAR
output_path = r'C:\Users\Bernardo\Documents\erp-ciete\erp-ciete\erp-ciete\Documentacion_ERP_Ciete.docx'
doc.save(output_path)
print(f'[OK] Documento guardado en: {output_path}')
